# =========================
# THƯ VIỆN BẮT BUỘC VÀ BỔ SUNG
# =========================
from datetime import datetime
import os
import numpy as np
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import (
    confusion_matrix,
    f1_score,
    accuracy_score,
    recall_score,
    precision_score,
    roc_auc_score,
    ConfusionMatrixDisplay,
)
import time

# Thư viện RSS Feed
try:
    import feedparser
    _FEEDPARSER_OK = True
except Exception:
    feedparser = None
    _FEEDPARSER_OK = False

# Thư viện GOOGLE GEMINI VÀ OPENAI (Giữ nguyên logic kiểm tra thư viện)
try:
    from google import genai
    from google.genai.errors import APIError
    _GEMINI_OK = True
except Exception:
    genai = None
    APIError = Exception
    _GEMINI_OK = False

try:
    from openai import OpenAI
    _OPENAI_OK = True
except Exception:
    OpenAI = None
    _OPENAI_OK = False

# Thư viện Word Export
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from io import BytesIO
    _WORD_OK = True
except Exception:
    _WORD_OK = False

MODEL_NAME = "gemini-2.5-flash"

# =========================
# HÀM TẠO WORD REPORT
# =========================

def generate_word_report(ratios_display, pd_value, pd_label, ai_analysis, fig_bar, fig_radar, company_name="KHÁCH HÀNG DOANH NGHIỆP"):
    """
    Tạo báo cáo Word chuyên nghiệp từ kết quả phân tích tín dụng.

    Parameters:
    - ratios_display: DataFrame chứa 14 chỉ số tài chính (index = tên chỉ số, column = giá trị)
    - pd_value: Xác suất vỡ nợ (PD) dưới dạng số float (0-1) hoặc NaN
    - pd_label: Nhãn dự đoán ("Default" hoặc "Non-Default")
    - ai_analysis: Text phân tích từ AI
    - fig_bar: Matplotlib figure của bar chart
    - fig_radar: Matplotlib figure của radar chart
    - company_name: Tên công ty (mặc định)

    Returns:
    - BytesIO object chứa Word document
    """

    if not _WORD_OK:
        raise Exception("Thiếu thư viện python-docx. Vui lòng cài đặt: pip install python-docx Pillow")

    # Tạo document mới
    doc = Document()

    # Cấu hình margin cho document
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ===== 1. HEADER VỚI LOGO VÀ TIÊU ĐỀ =====
    # Thêm logo nếu có
    try:
        if os.path.exists("logo-agribank.jpg"):
            doc.add_picture("logo-agribank.jpg", width=Inches(2.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass

    # Tiêu đề chính
    title = doc.add_heading('BÁO CÁO ĐÁNH GIÁ RỦI RO TÍN DỤNG', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(194, 24, 91)  # #c2185b
    title_run.font.bold = True

    # Subtitle
    subtitle = doc.add_paragraph('Dự báo Xác suất Vỡ nợ KHDN (PD) & Phân tích AI Chuyên sâu')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(255, 107, 157)  # #ff6b9d
    subtitle_run.font.bold = True

    # Thông tin thời gian
    date_info = doc.add_paragraph(f"Ngày xuất báo cáo: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_info.runs[0]
    date_run.font.size = Pt(10)

    # Thông tin khách hàng
    company_info = doc.add_paragraph()
    company_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_run = company_info.add_run(f"Tên khách hàng: {company_name}")
    company_run.font.size = Pt(11)
    company_run.font.bold = True

    doc.add_paragraph()  # Spacer

    # ===== 2. KẾT QUẢ DỰ BÁO PD =====
    heading1 = doc.add_heading('1. KẾT QUẢ DỰ BÁO XÁC SUẤT VỠ NỢ (PD)', level=1)
    heading1_run = heading1.runs[0]
    heading1_run.font.color.rgb = RGBColor(255, 107, 157)  # #ff6b9d

    pd_para = doc.add_paragraph()
    if pd.notna(pd_value):
        pd_para.add_run(f"Xác suất Vỡ nợ (PD): ").bold = True
        pd_para.add_run(f"{pd_value:.2%}\n")
        pd_para.add_run("Phân loại: ").bold = True
        pd_para.add_run(f"{pd_label}\n")

        if "Default" in pd_label and "Non-Default" not in pd_label:
            risk_run = pd_para.add_run("⚠️ RỦI RO CAO - CẦN XEM XÉT KỸ LƯỠNG")
            risk_run.bold = True
            risk_run.font.color.rgb = RGBColor(220, 53, 69)  # Red
        else:
            safe_run = pd_para.add_run("✓ RỦI RO THẤP - KHẢ QUAN")
            safe_run.bold = True
            safe_run.font.color.rgb = RGBColor(40, 167, 69)  # Green
    else:
        pd_para.add_run("Xác suất Vỡ nợ (PD): ").bold = True
        pd_para.add_run("Không có dữ liệu")

    doc.add_paragraph()  # Spacer

    # ===== 3. BẢNG CHỈ SỐ TÀI CHÍNH =====
    heading2 = doc.add_heading('2. CHỈ SỐ TÀI CHÍNH CHI TIẾT', level=1)
    heading2_run = heading2.runs[0]
    heading2_run.font.color.rgb = RGBColor(255, 107, 157)  # #ff6b9d

    # Tạo bảng
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Chỉ số Tài chính'
    hdr_cells[1].text = 'Giá trị'

    # Style header
    for cell in hdr_cells:
        cell_para = cell.paragraphs[0]
        cell_run = cell_para.runs[0]
        cell_run.font.bold = True
        cell_run.font.size = Pt(11)
        cell_run.font.color.rgb = RGBColor(255, 255, 255)
        # Set background color
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'FF6B9D')  # Pink
        cell._element.get_or_add_tcPr().append(shading_elm)
        cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for idx, row in ratios_display.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        value = row['Giá trị']
        row_cells[1].text = f"{value:.4f}" if pd.notna(value) else "N/A"
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()  # Spacer

    # ===== 4. BIỂU ĐỒ VISUALIZATION =====
    doc.add_page_break()
    heading3 = doc.add_heading('3. TRỰC QUAN HÓA DỮ LIỆU', level=1)
    heading3_run = heading3.runs[0]
    heading3_run.font.color.rgb = RGBColor(255, 107, 157)  # #ff6b9d

    # Bar chart
    try:
        doc.add_heading('3.1. Biểu đồ Cột - Giá trị các Chỉ số', level=2)
        bar_buffer = BytesIO()
        fig_bar.savefig(bar_buffer, format='png', dpi=150, bbox_inches='tight')
        bar_buffer.seek(0)
        doc.add_picture(bar_buffer, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Spacer
    except Exception as e:
        doc.add_paragraph(f"Không thể tạo biểu đồ cột: {str(e)}")

    # Radar chart
    try:
        doc.add_heading('3.2. Biểu đồ Radar - Phân tích Đa chiều', level=2)
        radar_buffer = BytesIO()
        fig_radar.savefig(radar_buffer, format='png', dpi=150, bbox_inches='tight')
        radar_buffer.seek(0)
        doc.add_picture(radar_buffer, width=Inches(5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"Không thể tạo biểu đồ radar: {str(e)}")

    # ===== 5. PHÂN TÍCH AI =====
    doc.add_page_break()
    heading4 = doc.add_heading('4. PHÂN TÍCH AI & KHUYẾN NGHỊ TÍN DỤNG', level=1)
    heading4_run = heading4.runs[0]
    heading4_run.font.color.rgb = RGBColor(255, 107, 157)  # #ff6b9d

    if ai_analysis and ai_analysis.strip():
        # Chia thành các đoạn và thêm vào document
        analysis_paragraphs = ai_analysis.split('\n')
        for para_text in analysis_paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text)
                # Highlight keywords
                if "CHO VAY" in para_text and "KHÔNG CHO VAY" not in para_text:
                    for run in para.runs:
                        if "CHO VAY" in run.text:
                            run.font.color.rgb = RGBColor(40, 167, 69)  # Green
                            run.bold = True
                elif "KHÔNG CHO VAY" in para_text:
                    for run in para.runs:
                        if "KHÔNG CHO VAY" in run.text:
                            run.font.color.rgb = RGBColor(220, 53, 69)  # Red
                            run.bold = True
    else:
        doc.add_paragraph("Chưa có phân tích từ AI. Vui lòng click nút 'Yêu cầu AI Phân tích & Đề xuất' để nhận khuyến nghị.")

    # ===== 6. FOOTER =====
    doc.add_paragraph()
    footer = doc.add_paragraph(
        f"Báo cáo này được tạo tự động bởi Hệ thống Đánh giá Rủi ro Tín dụng - Powered by AI & Machine Learning\n"
        f"© {datetime.now().year} Credit Risk Assessment System | Version 2.0 Premium"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.runs[0]
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)  # Grey

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# =========================
# CẤU HÌNH TRANG (NÂNG CẤP GIAO DIỆN)
# =========================
st.set_page_config(
    page_title="Credit Risk PD & Gemini Analysis | Banking Suite",
    page_icon="🏦",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ========================================
# CSS NÂNG CẤP - PHONG CÁCH NGÂN HÀNG HIỆN ĐẠI
# ========================================
st.markdown("""
<style>
/* ========== IMPORT GOOGLE FONTS ========== */
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700;900&family=Playfair+Display:wght@700;900&display=swap');

/* ========== GENERAL SETTINGS ========== */
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}

* {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
}

/* Main content area */
.main {
    background: linear-gradient(135deg, #fff5f7 0%, #ffe8f0 100%);
    animation: fadeIn 0.8s ease-in;
}

@keyframes fadeIn {
    from { opacity: 0; transform: translateY(20px); }
    to { opacity: 1; transform: translateY(0); }
}

/* ========== PREMIUM HEADER BANNER ========== */
.banner-title-container {
    background: linear-gradient(135deg, #ff6b9d 0%, #ff85a1 50%, #ff6b9d 100%);
    padding: 40px 50px;
    border-radius: 20px;
    box-shadow: 0 10px 40px rgba(255, 107, 157, 0.3),
                0 5px 15px rgba(255, 133, 161, 0.2),
                inset 0 1px 0 rgba(255, 255, 255, 0.1);
    margin-bottom: 30px;
    text-align: center;
    position: relative;
    overflow: hidden;
}

/* Shine effect */
.banner-title-container::before {
    content: '';
    position: absolute;
    top: -50%;
    left: -50%;
    width: 200%;
    height: 200%;
    background: linear-gradient(
        45deg,
        transparent 30%,
        rgba(255, 255, 255, 0.1) 50%,
        transparent 70%
    );
    animation: shine 3s infinite;
}

@keyframes shine {
    0% { transform: translateX(-100%) translateY(-100%) rotate(45deg); }
    100% { transform: translateX(100%) translateY(100%) rotate(45deg); }
}

.banner-title-container h1 {
    color: #ffffff !important;
    font-family: 'Playfair Display', serif !important;
    font-weight: 900 !important;
    font-size: 2.8rem !important;
    text-shadow: 2px 2px 8px rgba(0, 0, 0, 0.3),
                 0 0 30px rgba(255, 182, 193, 0.5);
    margin-bottom: 10px !important;
    letter-spacing: -0.5px;
    position: relative;
    z-index: 1;
    animation: titleGlow 2s ease-in-out infinite alternate;
}

@keyframes titleGlow {
    from { text-shadow: 2px 2px 8px rgba(0, 0, 0, 0.3), 0 0 30px rgba(255, 182, 193, 0.5); }
    to { text-shadow: 2px 2px 8px rgba(0, 0, 0, 0.3), 0 0 40px rgba(255, 182, 193, 0.7); }
}

.banner-title-container h3 {
    color: #fff0f5 !important;
    font-weight: 600 !important;
    font-size: 1.3rem !important;
    margin-top: 0 !important;
    border-bottom: none !important;
    letter-spacing: 1px;
    text-transform: uppercase;
    position: relative;
    z-index: 1;
}

/* Gold accent line */
.banner-title-container::after {
    content: '';
    position: absolute;
    bottom: 0;
    left: 50%;
    transform: translateX(-50%);
    width: 80%;
    height: 3px;
    background: linear-gradient(90deg, transparent, #ffb3c6, transparent);
    z-index: 1;
}

/* ========== SIDEBAR PREMIUM STYLING ========== */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #ff6b9d 0%, #e91e63 100%) !important;
    box-shadow: 2px 0 20px rgba(0, 0, 0, 0.1);
}

[data-testid="stSidebar"] * {
    color: #ffffff !important;
}

[data-testid="stSidebar"] .stMarkdown {
    color: #e8f4f8 !important;
}

/* File uploader trong sidebar */
div[data-testid="stFileUploader"] {
    background: rgba(255, 255, 255, 0.05);
    border: 2px dashed #ffb3c6 !important;
    border-radius: 15px;
    padding: 20px;
    margin-top: 15px;
    transition: all 0.3s ease;
    backdrop-filter: blur(10px);
}

div[data-testid="stFileUploader"]:hover {
    background: rgba(255, 255, 255, 0.1);
    border-color: #ffc0cb !important;
    transform: translateY(-2px);
    box-shadow: 0 5px 20px rgba(255, 179, 198, 0.3);
}

/* ========== TABS PREMIUM DESIGN ========== */
button[data-testid="stTab"] {
    background: linear-gradient(135deg, #ffffff 0%, #fff5f7 100%);
    border: 2px solid #ffd4dd;
    border-radius: 12px 12px 0 0 !important;
    transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
    font-weight: 700;
    font-size: 1rem;
    color: #4a5568;
    padding: 15px 30px;
    margin-right: 8px;
    box-shadow: 0 2px 8px rgba(0, 0, 0, 0.05);
}

button[data-testid="stTab"]:hover {
    background: linear-gradient(135deg, #ffe8f0 0%, #ffd4dd 100%);
    color: #c2185b;
    border-color: #ff6b9d;
    transform: translateY(-3px);
    box-shadow: 0 5px 15px rgba(255, 107, 157, 0.2);
}

button[data-testid="stTab"][aria-selected="true"] {
    background: linear-gradient(135deg, #ff6b9d 0%, #ff85a1 100%) !important;
    color: #ffffff !important;
    border-color: #ffb3c6 !important;
    border-bottom: 3px solid #ffb3c6 !important;
    box-shadow: 0 8px 20px rgba(255, 107, 157, 0.4),
                inset 0 1px 0 rgba(255, 255, 255, 0.2);
    transform: translateY(-3px);
}

/* ========== HEADINGS ========== */
h1, h2, h3, h4 {
    color: #1a2332 !important;
    font-weight: 700 !important;
}

h2 {
    color: #c2185b !important;
    border-bottom: 3px solid #ffb3c6;
    padding-bottom: 10px;
    margin-bottom: 20px !important;
}

h3 {
    color: #ff6b9d !important;
    border-bottom: 2px solid rgba(255, 179, 198, 0.3);
    padding-bottom: 8px;
    margin-bottom: 15px !important;
}

/* ========== METRIC CONTAINERS ========== */
div[data-testid="metric-container"] {
    background: linear-gradient(135deg, #ffffff 0%, #fff5f7 100%);
    border: 2px solid transparent;
    border-image: linear-gradient(135deg, #ffb3c6, #ff6b9d) 1;
    border-radius: 16px;
    padding: 20px;
    box-shadow: 0 8px 25px rgba(255, 107, 157, 0.12),
                0 3px 10px rgba(0, 0, 0, 0.08);
    transition: all 0.3s ease;
}

div[data-testid="metric-container"]:hover {
    transform: translateY(-5px) scale(1.02);
    box-shadow: 0 12px 35px rgba(255, 107, 157, 0.2),
                0 5px 15px rgba(255, 179, 198, 0.15);
}

/* Metric label */
div[data-testid="metric-container"] label {
    font-weight: 700 !important;
    color: #c2185b !important;
    font-size: 0.9rem !important;
    text-transform: uppercase;
    letter-spacing: 0.5px;
}

/* Metric value */
div[data-testid="metric-container"] [data-testid="stMetricValue"] {
    color: #ff6b9d !important;
    font-weight: 900 !important;
    font-size: 2.2rem !important;
}

/* ========== BUTTONS PREMIUM ========== */
button[kind="primary"] {
    background: linear-gradient(135deg, #ff6b9d 0%, #ff85a1 100%) !important;
    border: 2px solid #ffb3c6 !important;
    border-radius: 12px !important;
    color: #ffffff !important;
    font-weight: 700 !important;
    font-size: 1.05rem !important;
    padding: 12px 30px !important;
    transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1) !important;
    box-shadow: 0 6px 20px rgba(255, 107, 157, 0.3),
                inset 0 1px 0 rgba(255, 255, 255, 0.2);
    text-transform: uppercase;
    letter-spacing: 1px;
}

button[kind="primary"]:hover {
    background: linear-gradient(135deg, #e91e63 0%, #f06292 100%) !important;
    border-color: #ffc0cb !important;
    transform: translateY(-2px) scale(1.02);
    box-shadow: 0 10px 30px rgba(255, 107, 157, 0.4),
                0 5px 15px rgba(255, 179, 198, 0.3),
                inset 0 1px 0 rgba(255, 255, 255, 0.3);
}

button[kind="primary"]:active {
    transform: translateY(0) scale(0.98);
}

/* ========== CONTAINERS & CARDS ========== */
div[data-testid="stContainer"] {
    background: #ffffff;
    border-radius: 16px;
    padding: 25px;
    box-shadow: 0 4px 15px rgba(0, 0, 0, 0.08);
    border: 1px solid rgba(0, 61, 130, 0.1);
}

/* Expander */
div[data-testid="stExpander"] {
    background: #ffffff;
    border: 2px solid #ffd4dd;
    border-radius: 12px;
    box-shadow: 0 2px 8px rgba(0, 0, 0, 0.05);
    transition: all 0.3s ease;
}

div[data-testid="stExpander"]:hover {
    border-color: #ff6b9d;
    box-shadow: 0 4px 15px rgba(255, 107, 157, 0.15);
}

/* ========== DATAFRAMES ========== */
div[data-testid="stDataFrame"] {
    border: 2px solid #e0e6ed;
    border-radius: 12px;
    overflow: hidden;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
}

/* ========== INFO/WARNING/ERROR BOXES ========== */
div[data-baseweb="notification"] {
    border-radius: 12px;
    border-left-width: 5px !important;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
    padding: 20px !important;
}

/* Info box */
div[data-baseweb="notification"][data-testid*="stInfo"] {
    background: linear-gradient(135deg, #ffe8f0 0%, #ffd4dd 100%);
    border-left-color: #ff6b9d !important;
}

/* Success box */
div[data-baseweb="notification"][data-testid*="stSuccess"] {
    background: linear-gradient(135deg, #d4edda 0%, #c3e6cb 100%);
    border-left-color: #28a745 !important;
}

/* Warning box */
div[data-baseweb="notification"][data-testid*="stWarning"] {
    background: linear-gradient(135deg, #fff3cd 0%, #ffeaa7 100%);
    border-left-color: #ffc107 !important;
}

/* Error box */
div[data-baseweb="notification"][data-testid*="stError"] {
    background: linear-gradient(135deg, #f8d7da 0%, #f5c6cb 100%);
    border-left-color: #dc3545 !important;
}

/* ========== DIVIDER ========== */
hr {
    border: none;
    height: 3px;
    background: linear-gradient(90deg, transparent, #ffb3c6, transparent);
    margin: 30px 0;
}

/* ========== PROGRESS BAR ========== */
div[data-testid="stProgress"] > div {
    background: linear-gradient(90deg, #ff6b9d, #ff85a1, #ffb3c6);
    border-radius: 10px;
    box-shadow: 0 2px 10px rgba(255, 107, 157, 0.3);
}

/* ========== SPINNER ========== */
div[data-testid="stSpinner"] > div {
    border-top-color: #ffb3c6 !important;
}

/* ========== TOOLTIPS & CAPTIONS ========== */
.stCaption {
    color: #6b7280 !important;
    font-weight: 500 !important;
}

/* ========== RESPONSIVE ENHANCEMENTS ========== */
@media (max-width: 768px) {
    .banner-title-container {
        padding: 25px 20px;
    }

    .banner-title-container h1 {
        font-size: 2rem !important;
    }

    button[data-testid="stTab"] {
        padding: 10px 15px;
        font-size: 0.9rem;
    }
}

/* ========== SCROLL BAR ========== */
::-webkit-scrollbar {
    width: 10px;
    height: 10px;
}

::-webkit-scrollbar-track {
    background: #f5f7fa;
    border-radius: 10px;
}

::-webkit-scrollbar-thumb {
    background: linear-gradient(180deg, #ff6b9d, #ff85a1);
    border-radius: 10px;
}

::-webkit-scrollbar-thumb:hover {
    background: linear-gradient(180deg, #e91e63, #f06292);
}

</style>
""", unsafe_allow_html=True)


# =========================
# HÀM GỌI GEMINI API (GIỮ NGUYÊN LOGIC)
# =========================

def get_ai_analysis(data_payload: dict, api_key: str) -> str:
    """
    Sử dụng Gemini API để phân tích chỉ số tài chính.
    """
    if not _GEMINI_OK:
        return "Lỗi: Thiếu thư viện google-genai (cần cài đặt: pip install google-genai)."

    client = genai.Client(api_key=api_key)

    sys_prompt = (
        "Bạn là chuyên gia phân tích tín dụng doanh nghiệp tại ngân hàng Việt Nam. "
        "Phân tích toàn diện dựa trên 14 chỉ số tài chính được cung cấp và PD (Nếu có). Lưu ý PD trong mô hình này được tính theo bối cảnh doanh nghiệp Việt Nam"
        "Nêu rõ: (1) Khả năng sinh lời, (2) Thanh khoản, (3) Cơ cấu nợ, (4) Hiệu quả hoạt động. "
        "Kết thúc bằng khuyến nghị in hoa: CHO VAY hoặc KHÔNG CHO VAY, kèm 2–3 điều kiện nếu CHO VAY. "
        "Viết bằng tiếng Việt súc tích, chuyên nghiệp."
    )

    # Gửi tên tiếng Việt dễ hiểu hơn cho AI
    user_prompt = "Bộ chỉ số tài chính và PD cần phân tích:\n" + str(data_payload) + "\n\nHãy phân tích và đưa ra khuyến nghị."

    try:
        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=[
                {"role": "user", "parts": [{"text": sys_prompt + "\n\n" + user_prompt}]}
            ],
            config={"system_instruction": sys_prompt}
        )
        return response.text
    except APIError as e:
        return f"Lỗi gọi API Gemini: {e}"
    except Exception as e:
        return f"Lỗi không xác định: {e}"


def chat_with_gemini(user_message: str, api_key: str, context_data: dict = None) -> str:
    """
    Chatbot với Gemini AI để trả lời câu hỏi của người dùng về phân tích tín dụng.

    Args:
        user_message: Câu hỏi từ người dùng
        api_key: API key của Gemini
        context_data: Dữ liệu ngữ cảnh (chỉ số tài chính, PD, phân tích trước đó)

    Returns:
        Câu trả lời từ Gemini AI
    """
    if not _GEMINI_OK:
        return "Lỗi: Thiếu thư viện google-genai (cần cài đặt: pip install google-genai)."

    client = genai.Client(api_key=api_key)

    # System prompt cho chatbot
    sys_prompt = (
        "Bạn là chuyên gia tư vấn tín dụng doanh nghiệp tại ngân hàng. "
        "Nhiệm vụ của bạn là trả lời các câu hỏi của người dùng về phân tích tín dụng một cách chuyên nghiệp, "
        "dựa trên dữ liệu tài chính và phân tích đã được cung cấp. "
        "Trả lời súc tích, rõ ràng, dễ hiểu bằng tiếng Việt. "
        "Nếu cần, đưa ra các khuyến nghị hoặc giải thích chi tiết về các chỉ số tài chính."
    )

    # Tạo context prompt nếu có dữ liệu
    context_prompt = ""
    if context_data:
        context_prompt = "\n\nDữ liệu ngữ cảnh:\n" + str(context_data)

    full_prompt = user_message + context_prompt

    try:
        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=[
                {"role": "user", "parts": [{"text": full_prompt}]}
            ],
            config={"system_instruction": sys_prompt}
        )
        return response.text
    except APIError as e:
        return f"Lỗi gọi API Gemini: {e}"
    except Exception as e:
        return f"Lỗi không xác định: {e}"


# =========================
# HÀM LẤY DỮ LIỆU TÀI CHÍNH TỰ ĐỘNG TỪ GEMINI API
# =========================

@st.cache_data(ttl=2592000)  # Cache 30 ngày (tự động cập nhật mỗi tháng)
def get_financial_data_from_ai(api_key: str) -> pd.DataFrame:
    """
    Tự động lấy dữ liệu tài chính doanh nghiệp Việt Nam từ Gemini API.
    Dữ liệu bao gồm: Doanh thu, Tổng tài sản, Lợi nhuận, Nợ phải trả, VCSH theo quý.

    Returns:
        pd.DataFrame: DataFrame chứa dữ liệu tài chính theo quý
    """
    if not _GEMINI_OK:
        return None

    try:
        client = genai.Client(api_key=api_key)

        # Lấy quý hiện tại
        current_date = datetime.now()
        current_year = current_date.year
        current_month = current_date.month
        current_quarter = (current_month - 1) // 3 + 1

        # Prompt yêu cầu Gemini cung cấp dữ liệu tài chính
        sys_prompt = """Bạn là chuyên gia kinh tế và dữ liệu thống kê về doanh nghiệp Việt Nam.
        Hãy cung cấp dữ liệu tài chính tổng hợp của khu vực doanh nghiệp Việt Nam theo quý,
        dựa trên các nguồn thống kê đáng tin cậy như GSO (Tổng cục Thống kê Việt Nam),
        Bộ Kế hoạch và Đầu tư, hoặc các báo cáo kinh tế vĩ mô.

        Trả về dữ liệu dưới dạng JSON với cấu trúc sau:
        {
            "quarters": ["Q1-2021", "Q2-2021", ...],
            "revenue": [số liệu doanh thu tỷ VNĐ, ...],
            "assets": [số liệu tổng tài sản tỷ VNĐ, ...],
            "profit": [số liệu lợi nhuận tỷ VNĐ, ...],
            "debt": [số liệu nợ phải trả tỷ VNĐ, ...],
            "equity": [số liệu VCSH tỷ VNĐ, ...]
        }

        Chỉ trả về JSON, không giải thích thêm."""

        user_prompt = f"""Hãy cung cấp dữ liệu tài chính tổng hợp của khu vực doanh nghiệp Việt Nam
        từ quý Q1-2021 đến quý Q{current_quarter}-{current_year}.

        Bao gồm các chỉ số:
        - Doanh thu (Revenue) - tổng doanh thu khu vực doanh nghiệp, đơn vị tỷ VNĐ
        - Tổng tài sản (Total Assets) - tổng tài sản khu vực doanh nghiệp, đơn vị tỷ VNĐ
        - Lợi nhuận (Profit) - lợi nhuận sau thuế, đơn vị tỷ VNĐ
        - Nợ phải trả (Debt) - tổng nợ phải trả, đơn vị tỷ VNĐ
        - Vốn chủ sở hữu (Equity/VCSH) - tổng VCSH, đơn vị tỷ VNĐ

        Dữ liệu phải phản ánh xu hướng tăng trưởng thực tế của nền kinh tế Việt Nam.
        Chỉ trả về JSON thuần, không markdown, không giải thích."""

        response = client.models.generate_content(
            model=MODEL_NAME,
            contents=[
                {"role": "user", "parts": [{"text": sys_prompt + "\n\n" + user_prompt}]}
            ],
            config={"system_instruction": sys_prompt}
        )

        # Parse JSON response
        import json
        import re

        response_text = response.text.strip()

        # Loại bỏ markdown code block nếu có
        if "```json" in response_text:
            response_text = re.search(r'```json\s*(\{.*?\})\s*```', response_text, re.DOTALL).group(1)
        elif "```" in response_text:
            response_text = re.search(r'```\s*(\{.*?\})\s*```', response_text, re.DOTALL).group(1)

        data = json.loads(response_text)

        # Tạo DataFrame
        df = pd.DataFrame({
            'Quý': data.get('quarters', []),
            'Doanh thu (tỷ VNĐ)': data.get('revenue', []),
            'Tổng tài sản (tỷ VNĐ)': data.get('assets', []),
            'Lợi nhuận (tỷ VNĐ)': data.get('profit', []),
            'Nợ phải trả (tỷ VNĐ)': data.get('debt', []),
            'VCSH (tỷ VNĐ)': data.get('equity', [])
        })

        return df

    except Exception as e:
        st.error(f"Lỗi khi lấy dữ liệu từ AI: {e}")
        return None


# =========================
# TÍNH X1..X14 TỪ 3 SHEET (CDKT/BCTN/LCTT) - SỬ DỤNG TÊN TIẾNG VIỆT (GIỮ NGUYÊN)
# =========================

# Bảng ánh xạ Tên chỉ số tiếng Việt
COMPUTED_COLS = [
    "Biên Lợi nhuận Gộp (X1)", "Biên Lợi nhuận Tr.Thuế (X2)", "ROA Tr.Thuế (X3)", 
    "ROE Tr.Thuế (X4)", "Tỷ lệ Nợ/TTS (X5)", "Tỷ lệ Nợ/VCSH (X6)", 
    "Thanh toán Hiện hành (X7)", "Thanh toán Nhanh (X8)", "Khả năng Trả lãi (X9)", 
    "Khả năng Trả nợ Gốc (X10)", "Tỷ lệ Tiền/VCSH (X11)", "Vòng quay HTK (X12)", 
    "Kỳ thu tiền BQ (X13)", "Hiệu suất Tài sản (X14)"
]

# Alias các dòng quan trọng trong từng sheet (GIỮ NGUYÊN)
ALIAS_IS = {
    "doanh_thu_thuan": ["Doanh thu thuần", "Doanh thu bán hàng", "Doanh thu thuần về bán hàng và cung cấp dịch vụ"],
    "gia_von": ["Giá vốn hàng bán"],
    "loi_nhuan_gop": ["Lợi nhuận gộp"],
    "chi_phi_lai_vay": ["Chi phí lãi vay", "Chi phí tài chính (trong đó: chi phí lãi vay)"],
    "loi_nhuan_truoc_thue": ["Tổng lợi nhuận kế toán trước thuế", "Lợi nhuận trước thuế", "Lợi nhuận trước thuế thu nhập DN"],
}
ALIAS_BS = {
    "tong_tai_san": ["Tổng tài sản"],
    "von_chu_so_huu": ["Vốn chủ sở hữu", "Vốn CSH"],
    "no_phai_tra": ["Nợ phải trả"],
    "tai_san_ngan_han": ["Tài sản ngắn hạn"],
    "no_ngan_han": ["Nợ ngắn hạn"],
    "hang_ton_kho": ["Hàng tồn kho"],
    "tien_tdt": ["Tiền và các khoản tương đương tiền", "Tiền và tương đương tiền"],
    "phai_thu_kh": ["Phải thu ngắn hạn của khách hàng", "Phải thu khách hàng"],
    "no_dai_han_den_han": ["Nợ dài hạn đến hạn trả", "Nợ dài hạn đến hạn"],
}
ALIAS_CF = {
    "khau_hao": ["Khấu hao TSCĐ", "Khấu hao", "Chi phí khấu hao"],
}

def _pick_year_cols(df: pd.DataFrame):
    """Chọn 2 cột năm gần nhất từ sheet (ưu tiên cột có nhãn là năm)."""
    numeric_years = []
    for c in df.columns[1:]:
        try:
            y = int(float(str(c).strip()))
            if 1990 <= y <= 2100:
                numeric_years.append((y, c))
        except Exception:
            continue
    if numeric_years:
        numeric_years.sort(key=lambda x: x[0])
        return numeric_years[-2][1], numeric_years[-1][1]
    # fallback: 2 cột cuối
    cols = df.columns[-2:]
    return cols[0], cols[1]

def _get_row_vals(df: pd.DataFrame, aliases: list[str]):
    """Tìm dòng theo alias. Trả về (prev, cur) theo 2 cột năm gần nhất."""
    label_col = df.columns[0]
    prev_col, cur_col = _pick_year_cols(df)
    mask = False
    for alias in aliases:
        mask = mask | df[label_col].astype(str).str.contains(alias, case=False, na=False)
    rows = df[mask]
    if rows.empty:
        return np.nan, np.nan
    row = rows.iloc[0]

    def to_num(x):
        try:
            # Xóa dấu phẩy, khoảng trắng
            return float(str(x).replace(",", "").replace(" ", ""))
        except Exception:
            return np.nan

    return to_num(row[prev_col]), to_num(row[cur_col])

def compute_ratios_from_three_sheets(xlsx_file) -> pd.DataFrame:
    """Đọc 3 sheet CDKT/BCTN/LCTT và tính X1..X14 theo yêu cầu."""
    bs = pd.read_excel(xlsx_file, sheet_name="CDKT", engine="openpyxl")
    is_ = pd.read_excel(xlsx_file, sheet_name="BCTN", engine="openpyxl")
    cf = pd.read_excel(xlsx_file, sheet_name="LCTT", engine="openpyxl")

    # ---- Tính toán các biến số tài chính (GIỮ NGUYÊN CÁCH TÍNH)
    DTT_prev, DTT_cur         = _get_row_vals(is_, ALIAS_IS["doanh_thu_thuan"])
    GVHB_prev, GVHB_cur = _get_row_vals(is_, ALIAS_IS["gia_von"])
    LNG_prev, LNG_cur         = _get_row_vals(is_, ALIAS_IS["loi_nhuan_gop"])
    LNTT_prev, LNTT_cur = _get_row_vals(is_, ALIAS_IS["loi_nhuan_truoc_thue"])
    LV_prev, LV_cur           = _get_row_vals(is_, ALIAS_IS["chi_phi_lai_vay"])
    TTS_prev, TTS_cur           = _get_row_vals(bs, ALIAS_BS["tong_tai_san"])
    VCSH_prev, VCSH_cur         = _get_row_vals(bs, ALIAS_BS["von_chu_so_huu"])
    NPT_prev, NPT_cur           = _get_row_vals(bs, ALIAS_BS["no_phai_tra"])
    TSNH_prev, TSNH_cur         = _get_row_vals(bs, ALIAS_BS["tai_san_ngan_han"])
    NNH_prev, NNH_cur           = _get_row_vals(bs, ALIAS_BS["no_ngan_han"])
    HTK_prev, HTK_cur           = _get_row_vals(bs, ALIAS_BS["hang_ton_kho"])
    Tien_prev, Tien_cur         = _get_row_vals(bs, ALIAS_BS["tien_tdt"])
    KPT_prev, KPT_cur           = _get_row_vals(bs, ALIAS_BS["phai_thu_kh"])
    NDH_prev, NDH_cur           = _get_row_vals(bs, ALIAS_BS["no_dai_han_den_han"])
    KH_prev, KH_cur = _get_row_vals(cf, ALIAS_CF["khau_hao"])

    if pd.notna(GVHB_cur): GVHB_cur = abs(GVHB_cur)
    if pd.notna(LV_cur):      LV_cur     = abs(LV_cur)
    if pd.notna(KH_cur):      KH_cur     = abs(KH_cur)

    def avg(a, b):
        if pd.isna(a) and pd.isna(b): return np.nan
        if pd.isna(a): return b
        if pd.isna(b): return a
        return (a + b) / 2.0
    TTS_avg    = avg(TTS_cur,    TTS_prev)
    VCSH_avg = avg(VCSH_cur, VCSH_prev)
    HTK_avg    = avg(HTK_cur,    HTK_prev)
    KPT_avg    = avg(KPT_cur,    KPT_prev)

    EBIT_cur = (LNTT_cur + LV_cur) if (pd.notna(LNTT_cur) and pd.notna(LV_cur)) else np.nan
    NDH_cur = 0.0 if pd.isna(NDH_cur) else NDH_cur

    def div(a, b):
        return np.nan if (b is None or pd.isna(b) or b == 0) else a / b

    # ==== TÍNH X1..X14 ==== (GIỮ NGUYÊN CÔNG THỨC)
    X1  = div(LNG_cur, DTT_cur)
    X2  = div(LNTT_cur, DTT_cur)
    X3  = div(LNTT_cur, TTS_avg)
    X4  = div(LNTT_cur, VCSH_avg)
    X5  = div(NPT_cur,  TTS_cur)
    X6  = div(NPT_cur,  VCSH_cur)
    X7  = div(TSNH_cur, NNH_cur)
    X8  = div((TSNH_cur - HTK_cur) if pd.notna(TSNH_cur) and pd.notna(HTK_cur) else np.nan, NNH_cur)
    X9  = div(EBIT_cur, LV_cur)
    X10 = div((EBIT_cur + (KH_cur if pd.notna(KH_cur) else 0.0)), (LV_cur + NDH_cur) if pd.notna(LV_cur) else np.nan)
    X11 = div(Tien_cur, VCSH_cur)
    X12 = div(GVHB_cur, HTK_avg)
    turnover = div(DTT_cur, KPT_avg)
    X13 = div(365.0, turnover) if pd.notna(turnover) and turnover != 0 else np.nan
    X14 = div(DTT_cur, TTS_avg)

    # Khởi tạo DataFrame với tên cột tiếng Việt mới
    ratios = pd.DataFrame([[X1, X2, X3, X4, X5, X6, X7, X8, X9, X10, X11, X12, X13, X14]],
                          columns=COMPUTED_COLS)
                          
    # Thêm cột X_1..X_14 ẩn để phục vụ việc dự báo mô hình
    ratios[[f"X_{i}" for i in range(1, 15)]] = ratios.values
    return ratios

# =========================
# HÀM ĐỌC RSS FEED
# =========================

@st.cache_data(ttl=7200)  # Cache 120 phút = 7200 giây
def fetch_rss_feed(url, source_name):
    """
    Đọc RSS feed từ URL và trả về 5 bài mới nhất.

    Parameters:
    - url: Đường dẫn RSS feed
    - source_name: Tên nguồn tin

    Returns:
    - List của dict chứa {title, link, published}
    """
    if not _FEEDPARSER_OK:
        return [{"title": "⚠️ Thiếu thư viện feedparser", "link": "#", "published": ""}]

    try:
        feed = feedparser.parse(url)
        articles = []

        # Lấy 5 bài mới nhất
        for entry in feed.entries[:5]:
            title = entry.get('title', 'Không có tiêu đề')
            link = entry.get('link', '#')

            # Xử lý thời gian
            published = entry.get('published', '')
            if not published:
                published = entry.get('updated', '')

            # Parse thời gian nếu có
            pub_time = ""
            if published:
                try:
                    from dateutil import parser as date_parser
                    dt = date_parser.parse(published)
                    pub_time = dt.strftime('%d/%m/%Y %H:%M')
                except:
                    pub_time = published

            articles.append({
                'title': title,
                'link': link,
                'published': pub_time
            })

        return articles if articles else [{"title": "Không có bài viết mới", "link": "#", "published": ""}]

    except Exception as e:
        return [{"title": f"⚠️ Lỗi khi đọc RSS: {str(e)[:50]}", "link": "#", "published": ""}]

# =========================
# UI & TRAIN MODEL
# =========================
np.random.seed(0)

# ========================================
# PREMIUM BANKING HEADER
# ========================================
st.markdown('<div class="banner-title-container">', unsafe_allow_html=True)

# Thêm logo nếu có (optional)
col_logo, col_title = st.columns([1, 5])
with col_logo:
    try:
        st.image("logo-agribank.jpg", width=120)
    except:
        st.markdown("🏦")

with col_title:
    st.markdown("""
        <h1 style='margin: 0; padding: 0;'>CHƯƠNG TRÌNH ĐÁNH GIÁ RỦI RO TÍN DỤNG</h1>
        <h3 style='margin: 5px 0 0 0;'>Dự báo Xác suất Vỡ nợ KHDN (PD) & Phân tích AI Chuyên sâu</h3>
    """, unsafe_allow_html=True)

st.markdown('</div>', unsafe_allow_html=True)

# Load dữ liệu huấn luyện (CSV có default, X_1..X_14) - Giữ nguyên logic load data
try:
    df = pd.read_csv('DATASET.csv', encoding='latin-1')
    # Tên cột cho việc huấn luyện (phải giữ nguyên X_1..X_14)
    MODEL_COLS = [f"X_{i}" for i in range(1, 15)]
except Exception:
    df = None

# DI CHUYỂN UPLOADER VỀ ĐẦU SIDEBAR (Không còn selectbox)
uploaded_file = st.sidebar.file_uploader("📂 Tải CSV Dữ liệu Huấn luyện", type=['csv'])
if uploaded_file is not None:
    df = pd.read_csv(uploaded_file, encoding='latin-1')
    MODEL_COLS = [f"X_{i}" for i in range(1, 15)]
    
# Định nghĩa các Tabs
# ------------------------------------------------------------------------------------------------
# THAY ĐỔI 4: Vị trí Tabs được giữ nguyên, CSS mới sẽ đảm bảo Tabs có màu
# Tab mới: Dashboard tài chính doanh nghiệp (GSO) và Tin tức tài chính
# ------------------------------------------------------------------------------------------------
tab_predict, tab_dashboard, tab_news, tab_authors, tab_build, tab_goal = st.tabs([
    "🚀 Sử dụng mô hình dự báo",
    "📊 Dashboard tài chính doanh nghiệp",
    "📰 Tin tức tài chính",
    "👥 Nhóm tác giả",
    "🛠️ Xây dựng mô hình",
    "🎯 Mục tiêu của mô hình"
])

# --- Logic xử lý khi chưa có data huấn luyện ---
if df is None:
    st.sidebar.info("💡 Hãy tải file CSV huấn luyện (có cột 'default' và X_1...X_14) để xây dựng mô hình.")
    
    # Logic cho các tab khi thiếu data huấn luyện
    with tab_predict:
        st.header("⚡ Dự báo PD & Phân tích AI cho Hồ sơ mới")
        st.warning("⚠️ **Không thể dự báo PD**. Vui lòng tải file **CSV Dữ liệu Huấn luyện** ở sidebar để xây dựng mô hình Logistic Regression.")
        up_xlsx = st.file_uploader("Tải **ho_so_dn.xlsx**", type=["xlsx"], key="ho_so_dn")
        if up_xlsx is None:
            st.info("Hãy tải **ho_so_dn.xlsx** (đủ 3 sheet) để tính X1…X14 và phân tích AI.")

    with tab_goal:
        st.header("🎯 Mục tiêu của Mô hình")
        st.info("Ứng dụng này cần dữ liệu huấn luyện để bắt đầu hoạt động.")
    
    with tab_build:
          st.header("🛠️ Xây dựng & Đánh giá Mô hình LogReg")
          st.error("❌ **Không thể xây dựng mô hình**. Vui lòng tải file **CSV Dữ liệu Huấn luyện** ở sidebar để bắt đầu.")
          
    st.stop()
# ------------------------------------------------------------------------------------------------

# Hiển thị trạng thái thư viện AI (Sử dụng cột để bố trí đẹp hơn)
col_ai_status, col_date = st.columns([3, 1])
with col_ai_status:
    ai_status = ("✅ sẵn sàng (cần 'GEMINI_API_KEY' trong Secrets)" if _GEMINI_OK else "⚠️ Thiếu thư viện google-genai.")
    st.caption(f"🔎 Trạng thái Gemini AI: **<span style='color: #004c99; font-weight: bold;'>{ai_status}</span>**", unsafe_allow_html=True)
with col_date:
    st.caption(f"📅 Cập nhật: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

st.divider()

# Kiểm tra cột cần thiết
required_cols = ['default'] + MODEL_COLS
missing = [c for c in required_cols if c not in df.columns]
if missing:
    st.error(f"❌ Thiếu cột: **{missing}**. Vui lòng kiểm tra lại file CSV huấn luyện.")
    st.stop()


# Train model (GIỮ NGUYÊN)
X = df[MODEL_COLS] # Chỉ lấy các cột X_1..X_14
y = df['default'].astype(int)

X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42, stratify=y
)
model = LogisticRegression(random_state=42, max_iter=1000, class_weight="balanced", solver="lbfgs")
model.fit(X_train, y_train)

# Dự báo & đánh giá (GIỮ NGUYÊN)
y_pred_in = model.predict(X_train)
y_proba_in = model.predict_proba(X_train)[:, 1]
y_pred_out = model.predict(X_test)
y_proba_out = model.predict_proba(X_test)[:, 1]

metrics_in = {
    "accuracy_in": accuracy_score(y_train, y_pred_in),
    "precision_in": precision_score(y_train, y_pred_in, zero_division=0),
    "recall_in": recall_score(y_train, y_pred_in, zero_division=0),
    "f1_in": f1_score(y_train, y_pred_in, zero_division=0),
    "auc_in": roc_auc_score(y_train, y_proba_in),
}
metrics_out = {
    "accuracy_out": accuracy_score(y_test, y_pred_out),
    "precision_out": precision_score(y_test, y_pred_out, zero_division=0),
    "recall_out": recall_score(y_test, y_pred_out, zero_division=0),
    "f1_out": f1_score(y_test, y_pred_out, zero_division=0),
    "auc_out": roc_auc_score(y_test, y_proba_out),
}

# --- CÁC PHẦN UI DỰA TRÊN TABS ---

with tab_goal:
    st.header("🎯 Mục tiêu của Mô hình")
    st.markdown("**Dự báo xác suất vỡ nợ (PD) của khách hàng doanh nghiệp** dựa trên bộ chỉ số $\text{X1}–\text{X14}$ (tính từ Bảng Cân đối Kế toán, Báo cáo Kết quả Kinh doanh và Báo cáo Lưu chuyển Tiền tệ).")
    
    with st.expander("🖼️ Mô tả trực quan mô hình"):
        st.markdown("### Các hình ảnh minh họa cho mô hình Hồi quy Logistic và quy trình đánh giá rủi ro")

        # Hiển thị hình ảnh trong columns để layout đẹp hơn
        col_img1, col_img2 = st.columns(2)

        for idx, img in enumerate(["hinh2.jpg", "LogReg_1.png", "hinh3.png"]):
            try:
                if idx == 0:
                    with col_img1:
                        st.image(img, caption=f"Mô tả {idx+1}: Quy trình đánh giá", use_container_width=True)
                elif idx == 1:
                    with col_img2:
                        st.image(img, caption=f"Mô tả {idx+1}: Mô hình Logistic Regression", use_container_width=True)
                else:
                    st.image(img, caption=f"Mô tả {idx+1}: Kết quả phân tích", use_container_width=True)
            except Exception:
                # Nếu không tìm thấy file, hiển thị message thân thiện
                st.info(f"📊 Hình ảnh minh họa '{img}' sẽ được hiển thị ở đây")

    # Nút lên đầu trang
    st.markdown("""
        <div style='text-align: center; margin-top: 40px; margin-bottom: 20px;'>
            <a href='#top' onclick='window.scrollTo({top: 0, behavior: "smooth"}); return false;' style='text-decoration: none;'>
                <button style='
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    color: white;
                    border: none;
                    padding: 12px 30px;
                    border-radius: 25px;
                    font-size: 16px;
                    font-weight: 600;
                    cursor: pointer;
                    box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
                    transition: all 0.3s ease;
                '>
                    ⬆️ Lên đầu trang
                </button>
            </a>
        </div>
    """, unsafe_allow_html=True)

with tab_build:
    st.header("🛠️ Xây dựng & Đánh giá Mô hình LogReg")
    st.info("Mô hình Hồi quy Logistic đã được huấn luyện trên **20% dữ liệu Test (chưa thấy)**.")
    
    # Hiển thị Metrics quan trọng bằng st.metric
    st.subheader("1. Tổng quan Kết quả Đánh giá (Test Set)")
    col_acc, col_auc, col_f1 = st.columns(3)
    
    col_acc.metric(label="Độ chính xác (Accuracy)", value=f"{metrics_out['accuracy_out']:.2%}")
    # Đảm bảo logic delta vẫn đúng
    col_auc.metric(label="Diện tích dưới đường cong (AUC)", value=f"{metrics_out['auc_out']:.3f}", delta=f"{metrics_in['auc_in'] - metrics_out['auc_out']:.3f}", delta_color="inverse")
    col_f1.metric(label="Điểm F1-Score", value=f"{metrics_out['f1_out']:.3f}")
    
    st.divider()

    # Thống kê chi tiết & Biểu đồ
    st.subheader("2. Dữ liệu và Trực quan hóa")
    
    with st.expander("📊 Thống kê Mô tả và Dữ liệu Mẫu"):
        st.markdown("##### Thống kê Mô tả các biến $X_1..X_{14}$")
        st.dataframe(df[MODEL_COLS].describe().style.format("{:.4f}"))
        st.markdown("##### 6 Dòng dữ liệu huấn luyện mẫu (Đầu/Cuối)")
        st.dataframe(pd.concat([df.head(3), df.tail(3)]))

    st.markdown("##### Biểu đồ Phân tán (Scatter Plot) với Đường Hồi quy Logisitc")
    col = st.selectbox('🔍 Chọn biến X muốn vẽ', options=MODEL_COLS, index=0, key="select_build_col")
    
    # Biểu đồ Scatter Plot và Đường Hồi quy Logisitc (GIỮ NGUYÊN LOGIC, CẢI THIỆN MÀU SẮC)
    if col in df.columns:
        try:
            # Dùng Streamlit.pyplot với theme banking hiện đại
            fig, ax = plt.subplots(figsize=(12, 7))

            # Set background color
            fig.patch.set_facecolor('#f8f9fa')
            ax.set_facecolor('#ffffff')

            # Scatter plot với màu sắc pink rose theme
            sns.scatterplot(data=df, x=col, y='default', alpha=0.65, ax=ax, hue='default',
                          palette=['#ff6b9d', '#ffb3c6'], s=80, edgecolor='white', linewidth=0.5)

            # Vẽ đường logistic regression theo 1 biến
            x_range = np.linspace(df[col].min(), df[col].max(), 100).reshape(-1, 1)
            X_temp = df[[col]].copy()
            y_temp = df['default']
            lr_temp = LogisticRegression(max_iter=1000)
            lr_temp.fit(X_temp, y_temp)
            x_test = pd.DataFrame({col: x_range[:, 0]})
            y_curve = lr_temp.predict_proba(x_test)[:, 1]
            ax.plot(x_range, y_curve, color='#c2185b', linewidth=4, label='Đường LogReg',
                   linestyle='-', alpha=0.9)

            # Styling cho tiêu đề và labels
            ax.set_title(f'Quan hệ giữa {col} và Xác suất Vỡ nợ', fontsize=16, fontweight='bold', color='#c2185b', pad=20)
            ax.set_ylabel('Xác suất Default (0: Non-Default, 1: Default)', fontsize=13, fontweight='600', color='#4a5568')
            ax.set_xlabel(col, fontsize=13, fontweight='600', color='#4a5568')

            # Grid styling
            ax.grid(True, alpha=0.2, linestyle='--', linewidth=0.8, color='#ff6b9d')
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            ax.spines['left'].set_color('#d0d0d0')
            ax.spines['bottom'].set_color('#d0d0d0')

            # Legend styling
            legend = ax.legend(title='Default Status', title_fontsize=11, fontsize=10,
                             frameon=True, fancybox=True, shadow=True)
            legend.get_frame().set_facecolor('#f8f9fa')
            legend.get_frame().set_alpha(0.9)

            st.pyplot(fig)
            plt.close(fig)
        except Exception as e:
            st.error(f"Lỗi khi vẽ biểu đồ: {e}")
    else:
        st.warning("Biến không tồn tại trong dữ liệu.")
    
    st.divider()

    st.subheader("3. Ma trận Nhầm lẫn và Bảng Metrics Chi tiết")
    col_cm, col_metrics_table = st.columns(2)
    
    with col_cm:
        st.markdown("##### Ma trận Nhầm lẫn (Test Set)")
        cm = confusion_matrix(y_test, y_pred_out)

        # Tạo custom colormap cho pink rose theme
        from matplotlib.colors import LinearSegmentedColormap
        colors_pink = ['#fff5f7', '#ffe8f0', '#ffd4dd', '#ff85a1', '#ff6b9d']
        n_bins = 100
        cmap_pink = LinearSegmentedColormap.from_list('pink_rose', colors_pink, N=n_bins)

        disp = ConfusionMatrixDisplay(confusion_matrix=cm, display_labels=['Non-Default (0)', 'Default (1)'])
        fig2, ax = plt.subplots(figsize=(7, 7))
        fig2.patch.set_facecolor('#f8f9fa')

        disp.plot(ax=ax, cmap=cmap_pink, colorbar=True)

        # Styling
        ax.set_title('Ma trận Nhầm lẫn', fontsize=14, fontweight='bold', color='#c2185b', pad=15)
        ax.set_xlabel('Predicted Label', fontsize=12, fontweight='600', color='#4a5568')
        ax.set_ylabel('True Label', fontsize=12, fontweight='600', color='#4a5568')

        st.pyplot(fig2)
        plt.close(fig2)
        
    with col_metrics_table:
        st.markdown("##### Bảng Metrics Chi tiết")
        dt = pd.DataFrame({
            "Metric": ["Accuracy", "Precision", "Recall", "F1-Score", "AUC"],
            "Train Set": [metrics_in['accuracy_in'], metrics_in['precision_in'], metrics_in['recall_in'], metrics_in['f1_in'], metrics_in['auc_in']],
            "Test Set": [metrics_out['accuracy_out'], metrics_out['precision_out'], metrics_out['recall_out'], metrics_out['f1_out'], metrics_out['auc_out']],
        }).set_index("Metric")
        # Thêm styling để làm nổi bật kết quả tốt nhất
        def highlight_max(s):
            is_max = s == s.max()
            return ['background-color: #e0f0ff' if v else '' for v in is_max]

        st.dataframe(dt.style.format("{:.4f}").apply(highlight_max, axis=1), use_container_width=True)

    # Nút lên đầu trang
    st.markdown("""
        <div style='text-align: center; margin-top: 40px; margin-bottom: 20px;'>
            <a href='#top' onclick='window.scrollTo({top: 0, behavior: "smooth"}); return false;' style='text-decoration: none;'>
                <button style='
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    color: white;
                    border: none;
                    padding: 12px 30px;
                    border-radius: 25px;
                    font-size: 16px;
                    font-weight: 600;
                    cursor: pointer;
                    box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
                    transition: all 0.3s ease;
                '>
                    ⬆️ Lên đầu trang
                </button>
            </a>
        </div>
    """, unsafe_allow_html=True)

with tab_predict:
    # Trang này được hiển thị mặc định
    st.header("⚡ Dự báo PD & Phân tích AI cho Hồ sơ mới")
    
    # Sử dụng st.container và st.expander để tổ chức khu vực upload
    input_container = st.container(border=True)
    with input_container:
        st.markdown("##### 📥 Tải lên Hồ sơ Doanh nghiệp (Excel)")
        st.caption("File phải có đủ **3 sheet**: **CDKT** (Bảng Cân đối Kế toán) ; **BCTN** (Báo cáo Kết quả Kinh doanh) ; **LCTT** (Báo cáo Lưu chuyển Tiền tệ).")
        up_xlsx = st.file_uploader("Tải **ho_so_dn.xlsx**", type=["xlsx"], key="ho_so_dn_main", label_visibility="collapsed")
    
    if up_xlsx is not None:
        # Tính X1..X14 từ 3 sheet (GIỮ NGUYÊN)
        try:
            # Hiển thị thanh tiến trình giả lập (thêm hiệu ứng động)
            with st.spinner('Đang đọc và xử lý dữ liệu tài chính...'):
                ratios_df = compute_ratios_from_three_sheets(up_xlsx)
            
            # Tách riêng 14 cột tiếng Việt (hiển thị) và 14 cột tiếng Anh (dự báo)
            # ratios_display là DataFrame 1 cột: Index (Tên chỉ số) | Giá trị
            ratios_display = ratios_df[COMPUTED_COLS].T.rename(columns={0: 'Giá trị'})
            ratios_predict = ratios_df[MODEL_COLS]
            
        except Exception as e:
            st.error(f"❌ Lỗi tính chỉ số tài chính: Vui lòng kiểm tra lại cấu trúc 3 sheet trong file Excel. Chi tiết lỗi: {e}")
            st.stop()

        st.divider()
        st.markdown("### 1. 🔢 Các Chỉ số Tài chính Đã tính")
        
        # Tạo payload data cho AI (Sử dụng tên tiếng Việt)
        data_for_ai = ratios_display.to_dict()['Giá trị']
        
        # (Tuỳ chọn) dự báo PD nếu mô hình đã huấn luyện đúng cấu trúc X_1..X_14
        probs = np.nan
        preds = np.nan
        # Kiểm tra mô hình có sẵn sàng dự báo không (đã train và cột khớp)
        if set(X.columns) == set(ratios_predict.columns):
            try:
                # Đảm bảo thứ tự cột cho predict đúng như thứ tự cột huấn luyện
                probs_array = model.predict_proba(ratios_predict[X.columns])[:, 1]
                # Chuyển từ numpy array sang scalar để tránh lỗi ambiguous truth value
                probs = float(probs_array[0])
                preds = int(probs >= 0.15)
                # Thêm PD vào payload AI
                data_for_ai['Xác suất Vỡ nợ (PD)'] = probs
                data_for_ai['Dự đoán PD'] = "Default (Vỡ nợ)" if preds == 1 else "Non-Default (Không vỡ nợ)"
            except Exception as e:
                # Nếu có lỗi dự báo, chỉ cảnh báo, không dừng app
                st.warning(f"Không dự báo được PD: {e}")
        
        # ------------------------------------------------------------------------------------------------
        # ĐIỀU CHỈNH CỦA CHUYÊN GIA PYTHON: Bỏ .T để hiển thị đúng Tên Biến | Con số
        # ------------------------------------------------------------------------------------------------
        pd_col_1, pd_col_2, pd_col_pd = st.columns([2, 2, 1]) # Chia làm 3 cột, 2 cột giữa hiển thị ratios, 1 cột cuối hiển thị PD
        
        ratios_list = ratios_display.index.tolist()
        mid_point = len(ratios_list) // 2
        # ratios_display đã có cấu trúc đúng: Index (Tên biến) | Giá trị (Con số)
        ratios_part1 = ratios_display.iloc[:mid_point]
        ratios_part2 = ratios_display.iloc[mid_point:]
        
        # Hàm styling (GIỮ NGUYÊN)
        def color_ratios(val):
            """Ánh xạ màu dựa trên tên chỉ số và giá trị (tạm thời để hiển thị đẹp)"""
            # Chỉ số Thanh khoản (X7, X8) - Green/Yellow
            if "Thanh toán" in val.name and val.values[0] < 1.0: return ['background-color: #ffcccc' for _ in val] # Dưới 1: Báo động đỏ
            if "Thanh toán" in val.name and val.values[0] > 1.5: return ['background-color: #ccffcc' for _ in val] # Trên 1.5: Tốt
            # Chỉ số Nợ (X5, X6) - Red/Green
            if "Tỷ lệ Nợ/" in val.name and val.values[0] > 1.0: return ['background-color: #ffcccc' for _ in val] # Trên 1: Rủi ro cao
            if "Tỷ lệ Nợ/" in val.name and val.values[0] < 0.5: return ['background-color: #ccffcc' for _ in val] # Dưới 0.5: Tốt
            # Chỉ số Sinh lời (X1, X2, X3, X4) - Green/Yellow
            if "Lợi nhuận" in val.name or "ROA" in val.name or "ROE" in val.name:
                if val.values[0] <= 0: return ['background-color: #ffcccc' for _ in val]
                if val.values[0] > 0.1: return ['background-color: #ccffcc' for _ in val]
            return [''] * len(val)

        with pd_col_1:
             # Đảm bảo hiển thị Tên biến | Giá trị
             st.markdown("##### **Chỉ số Tài chính (1/2)**") 
             st.dataframe(
                 ratios_part1.style.apply(color_ratios, axis=1).format("{:.4f}").set_properties(**{'font-size': '14px'}),
                 use_container_width=True
             )

        with pd_col_2:
            # Đảm bảo hiển thị Tên biến | Giá trị
            st.markdown("##### **Chỉ số Tài chính (2/2)**")
            st.dataframe(
                ratios_part2.style.apply(color_ratios, axis=1).format("{:.4f}").set_properties(**{'font-size': '14px'}),
                use_container_width=True
            )
        
        with pd_col_pd:
            pd_value = f"{probs:.2%}" if pd.notna(probs) else "N/A"
            pd_delta = "⬆️ Rủi ro cao" if pd.notna(preds) and preds == 1 else "⬇️ Rủi ro thấp"

            st.metric(
                label="**Xác suất Vỡ nợ (PD)**",
                value=pd_value,
                delta=pd_delta if pd.notna(probs) else None,
                # Đảo ngược màu sắc delta cho PD: Rủi ro cao là màu đỏ (inverse), rủi ro thấp là màu xanh (normal)
                delta_color=("inverse" if pd.notna(preds) and preds == 1 else "normal")
            )
        # ------------------------------------------------------------------------------------------------

        st.divider()

        # ========================================
        # THÊM BIỂU ĐỒ VISUALIZATION CHO CÁC CHỈ SỐ TÀI CHÍNH
        # ========================================
        st.markdown("### 2. 📊 Trực quan hóa Các Chỉ số Tài chính")

        # Tạo 2 cột cho 2 loại biểu đồ
        chart_col1, chart_col2 = st.columns(2)

        with chart_col1:
            st.markdown("#### 📈 Biểu đồ Cột - Giá trị các Chỉ số")
            # Tạo bar chart
            fig_bar, ax_bar = plt.subplots(figsize=(8, 10))
            fig_bar.patch.set_facecolor('#fff5f7')
            ax_bar.set_facecolor('#ffffff')

            # Chuẩn bị data cho bar chart
            indicators = ratios_display.index.tolist()
            values = ratios_display['Giá trị'].values

            # Tạo màu gradient cho các bars
            bar_colors = plt.cm.RdPu(np.linspace(0.3, 0.9, len(indicators)))

            # Vẽ horizontal bar chart
            bars = ax_bar.barh(indicators, values, color=bar_colors, edgecolor='white', linewidth=1.5)

            # Thêm giá trị vào cuối mỗi bar
            for i, (bar, val) in enumerate(zip(bars, values)):
                width = bar.get_width()
                ax_bar.text(width, bar.get_y() + bar.get_height()/2,
                           f' {val:.3f}', ha='left', va='center',
                           fontsize=9, fontweight='600', color='#c2185b')

            # Styling
            ax_bar.set_xlabel('Giá trị', fontsize=12, fontweight='600', color='#4a5568')
            ax_bar.set_title('Các Chỉ số Tài chính', fontsize=14, fontweight='bold', color='#c2185b', pad=15)
            ax_bar.grid(True, alpha=0.2, linestyle='--', linewidth=0.8, color='#ff6b9d', axis='x')
            ax_bar.spines['top'].set_visible(False)
            ax_bar.spines['right'].set_visible(False)
            ax_bar.spines['left'].set_color('#d0d0d0')
            ax_bar.spines['bottom'].set_color('#d0d0d0')

            # Điều chỉnh layout để labels không bị cắt
            plt.tight_layout()
            st.pyplot(fig_bar)
            plt.close(fig_bar)

        with chart_col2:
            st.markdown("#### 🎯 Biểu đồ Radar - Phân tích Đa chiều")
            # Tạo radar chart (spider chart)
            fig_radar = plt.figure(figsize=(10, 10))
            fig_radar.patch.set_facecolor('#fff5f7')
            ax_radar = fig_radar.add_subplot(111, projection='polar')

            # Chuẩn bị data cho radar chart
            # Normalize các giá trị về khoảng 0-1 để dễ visualize
            from sklearn.preprocessing import MinMaxScaler
            scaler = MinMaxScaler()
            normalized_values = scaler.fit_transform(values.reshape(-1, 1)).flatten()

            # Tạo các góc cho mỗi chỉ số
            angles = np.linspace(0, 2 * np.pi, len(indicators), endpoint=False).tolist()
            normalized_values = normalized_values.tolist()

            # Đóng vòng tròn
            angles += angles[:1]
            normalized_values += normalized_values[:1]

            # Vẽ radar chart
            ax_radar.plot(angles, normalized_values, 'o-', linewidth=2.5, color='#ff6b9d', label='Chỉ số')
            ax_radar.fill(angles, normalized_values, alpha=0.25, color='#ffb3c6')

            # Thêm labels
            ax_radar.set_xticks(angles[:-1])
            # Rút ngắn tên chỉ số để dễ đọc
            short_labels = [label.split('(')[0].strip()[:20] for label in indicators]
            ax_radar.set_xticklabels(short_labels, size=8, color='#4a5568', fontweight='600')

            # Styling
            ax_radar.set_ylim(0, 1)
            ax_radar.set_title('Phân tích Đa chiều các Chỉ số\n(Normalized 0-1)',
                              fontsize=14, fontweight='bold', color='#c2185b', pad=20)
            ax_radar.grid(True, alpha=0.3, linestyle='--', linewidth=0.8, color='#ff6b9d')
            ax_radar.set_facecolor('#ffffff')

            plt.tight_layout()
            st.pyplot(fig_radar)
            plt.close(fig_radar)

        # Thêm expander với thông tin bổ sung
        with st.expander("ℹ️ Giải thích về Biểu đồ"):
            st.markdown("""
            **Biểu đồ Cột (Bar Chart):**
            - Hiển thị giá trị thực tế của từng chỉ số tài chính
            - Màu sắc gradient từ nhạt đến đậm để dễ phân biệt
            - Giá trị cụ thể được hiển thị bên cạnh mỗi cột

            **Biểu đồ Radar (Spider Chart):**
            - Hiển thị cân bằng tổng thể giữa các chỉ số
            - Giá trị được chuẩn hóa về thang 0-1 để dễ so sánh
            - Diện tích vùng phủ thể hiện độ mạnh của các chỉ số
            - Hình dạng đều = tốt, hình dạng lệch = cần cân bằng
            """)

        st.divider()

        # Khu vực Phân tích AI
        st.markdown("### 3. 🧠 Phân tích AI & Khuyến nghị Tín dụng")

        # Khởi tạo session_state cho phân tích AI
        if 'show_ai_analysis' not in st.session_state:
            st.session_state['show_ai_analysis'] = False
        if 'ai_analysis' not in st.session_state:
            st.session_state['ai_analysis'] = ''
        if 'chat_messages' not in st.session_state:
            st.session_state['chat_messages'] = []
        if 'ai_context_data' not in st.session_state:
            st.session_state['ai_context_data'] = {}

        ai_container = st.container(border=True)
        with ai_container:
            st.markdown("Sử dụng AI để phân tích toàn diện các chỉ số và đưa ra khuyến nghị chuyên nghiệp.")

            # Tạo 2 cột cho nút phân tích và nút ẩn
            col_btn1, col_btn2 = st.columns([3, 1])

            with col_btn1:
                analyze_button = st.button("✨ Yêu cầu AI Phân tích & Đề xuất", use_container_width=True, type="primary", key="analyze_ai_btn")

            with col_btn2:
                if st.session_state['show_ai_analysis']:
                    hide_button = st.button("🔽 Ẩn phân tích", use_container_width=True, key="hide_ai_btn")
                    if hide_button:
                        st.session_state['show_ai_analysis'] = False
                        st.session_state['chat_messages'] = []
                        st.rerun()

            # Xử lý khi người dùng click nút phân tích
            if analyze_button:
                # Kiểm tra API Key: ưu tiên lấy từ secrets
                api_key = st.secrets.get("GEMINI_API_KEY")

                if api_key:
                    # Thêm thanh tiến trình đẹp mắt
                    progress_bar = st.progress(0, text="Đang gửi dữ liệu và chờ Gemini phân tích...")
                    for percent_complete in range(100):
                        import time
                        time.sleep(0.01) # Giả lập thời gian xử lý
                        progress_bar.progress(percent_complete + 1, text=f"Đang gửi dữ liệu và chờ Gemini phân tích... {percent_complete+1}%")

                    ai_result = get_ai_analysis(data_for_ai, api_key)
                    progress_bar.empty() # Xóa thanh tiến trình

                    # Lưu kết quả vào session_state
                    st.session_state['ai_analysis'] = ai_result
                    st.session_state['show_ai_analysis'] = True
                    st.session_state['ai_context_data'] = data_for_ai
                    st.session_state['chat_messages'] = []  # Reset chat khi phân tích mới
                    st.rerun()
                else:
                    st.error("❌ **Lỗi Khóa API**: Không tìm thấy Khóa API. Vui lòng cấu hình Khóa **'GEMINI_API_KEY'** trong Streamlit Secrets.")

        # Hiển thị kết quả phân tích AI và chatbot nếu đã có phân tích
        if st.session_state['show_ai_analysis'] and st.session_state['ai_analysis']:
            ai_result = st.session_state['ai_analysis']

            st.markdown("---")
            st.markdown("**Kết quả Phân tích Chi tiết từ Gemini AI:**")

            if "KHÔNG CHO VAY" in ai_result.upper():
                st.error("🚨 **KHUYẾN NGHỊ CUỐI CÙNG: KHÔNG CHO VAY**")
                st.snow()
            elif "CHO VAY" in ai_result.upper():
                st.success("✅ **KHUYẾN NGHỊ CUỐI CÙNG: CHO VAY**")
                st.balloons()
            else:
                st.info("💡 **KHUYẾN NGHỊ CUỐI CÙNG**")

            st.info(ai_result)

            # ===== CHATBOT GEMINI AI =====
            st.markdown("---")
            st.markdown("#### 💬 Chatbot - Hỏi thêm thông tin")

            # Container cho chatbot
            chatbot_container = st.container(border=True)
            with chatbot_container:
                st.markdown("Bạn có thể hỏi thêm về kết quả phân tích, các chỉ số tài chính, hoặc bất kỳ câu hỏi nào liên quan đến tín dụng.")

                # Hiển thị lịch sử chat
                if st.session_state['chat_messages']:
                    st.markdown("**Lịch sử trò chuyện:**")
                    for msg in st.session_state['chat_messages']:
                        if msg['role'] == 'user':
                            st.markdown(f"**👤 Bạn:** {msg['content']}")
                        else:
                            st.markdown(f"**🤖 Gemini AI:** {msg['content']}")
                    st.markdown("---")

                # Form nhập câu hỏi
                with st.form(key='chat_form', clear_on_submit=True):
                    user_question = st.text_input(
                        "Nhập câu hỏi của bạn:",
                        placeholder="VD: Giải thích thêm về chỉ số thanh khoản...",
                        key='user_question_input'
                    )

                    col1, col2 = st.columns([1, 5])
                    with col1:
                        submit_button = st.form_submit_button("📤 Gửi", use_container_width=True)
                    with col2:
                        clear_button = st.form_submit_button("🗑️ Xóa lịch sử chat", use_container_width=True)

                # Xử lý khi người dùng gửi câu hỏi
                if submit_button and user_question.strip():
                    # Lấy API key
                    api_key = st.secrets.get("GEMINI_API_KEY")

                    # Lưu câu hỏi của user
                    st.session_state['chat_messages'].append({
                        'role': 'user',
                        'content': user_question
                    })

                    # Chuẩn bị context data cho chatbot
                    context_data = {
                        'chỉ_số_tài_chính': st.session_state.get('ai_context_data', data_for_ai),
                        'phân_tích_trước_đó': st.session_state['ai_analysis']
                    }

                    # Gọi chatbot API
                    with st.spinner("🤔 Gemini đang suy nghĩ..."):
                        bot_response = chat_with_gemini(user_question, api_key, context_data)

                    # Lưu response của bot
                    st.session_state['chat_messages'].append({
                        'role': 'assistant',
                        'content': bot_response
                    })

                    # Rerun để hiển thị tin nhắn mới
                    st.rerun()

                # Xử lý khi người dùng xóa lịch sử
                if clear_button:
                    st.session_state['chat_messages'] = []
                    st.rerun()

        st.divider()

        # ===== NÚT XUẤT FILE WORD =====
        st.markdown("### 4. 📄 Xuất Báo cáo Word")

        export_container = st.container(border=True)
        with export_container:
            st.markdown("Xuất toàn bộ phân tích (chỉ số tài chính, biểu đồ, PD, khuyến nghị AI) ra file Word chuyên nghiệp.")

            col_export1, col_export2 = st.columns([3, 1])

            with col_export1:
                company_name_input = st.text_input("Tên Khách hàng (tùy chọn):", value="KHÁCH HÀNG DOANH NGHIỆP", key="company_name_word")

            with col_export2:
                st.write("")  # Spacer

            if st.button("📥 Xuất file Word", use_container_width=True, type="primary", key="export_word_btn"):
                if not _WORD_OK:
                    st.error("❌ Thiếu thư viện python-docx. Không thể xuất Word.")
                else:
                    try:
                        with st.spinner("Đang tạo báo cáo Word..."):
                            # Lấy AI analysis từ session_state nếu có
                            ai_analysis_text = st.session_state.get('ai_analysis', '')

                            # Tạo lại figures để export (không hiển thị)
                            # Bar chart
                            fig_bar_export, ax_bar_export = plt.subplots(figsize=(8, 10))
                            fig_bar_export.patch.set_facecolor('#fff5f7')
                            ax_bar_export.set_facecolor('#ffffff')

                            indicators_export = ratios_display.index.tolist()
                            values_export = ratios_display['Giá trị'].values
                            bar_colors_export = plt.cm.RdPu(np.linspace(0.3, 0.9, len(indicators_export)))

                            bars_export = ax_bar_export.barh(indicators_export, values_export, color=bar_colors_export, edgecolor='white', linewidth=1.5)

                            for i, (bar, val) in enumerate(zip(bars_export, values_export)):
                                width = bar.get_width()
                                ax_bar_export.text(width, bar.get_y() + bar.get_height()/2,
                                           f' {val:.3f}', ha='left', va='center',
                                           fontsize=9, fontweight='600', color='#c2185b')

                            ax_bar_export.set_xlabel('Giá trị', fontsize=12, fontweight='600', color='#4a5568')
                            ax_bar_export.set_title('Các Chỉ số Tài chính', fontsize=14, fontweight='bold', color='#c2185b', pad=15)
                            ax_bar_export.grid(True, alpha=0.2, linestyle='--', linewidth=0.8, color='#ff6b9d', axis='x')
                            ax_bar_export.spines['top'].set_visible(False)
                            ax_bar_export.spines['right'].set_visible(False)
                            ax_bar_export.spines['left'].set_color('#d0d0d0')
                            ax_bar_export.spines['bottom'].set_color('#d0d0d0')
                            plt.tight_layout()

                            # Radar chart
                            fig_radar_export = plt.figure(figsize=(10, 10))
                            fig_radar_export.patch.set_facecolor('#fff5f7')
                            ax_radar_export = fig_radar_export.add_subplot(111, projection='polar')

                            from sklearn.preprocessing import MinMaxScaler
                            scaler_export = MinMaxScaler()
                            normalized_values_export = scaler_export.fit_transform(values_export.reshape(-1, 1)).flatten()

                            angles_export = np.linspace(0, 2 * np.pi, len(indicators_export), endpoint=False).tolist()
                            normalized_values_list_export = normalized_values_export.tolist()

                            angles_export += angles_export[:1]
                            normalized_values_list_export += normalized_values_list_export[:1]

                            ax_radar_export.plot(angles_export, normalized_values_list_export, 'o-', linewidth=2.5, color='#ff6b9d', label='Chỉ số')
                            ax_radar_export.fill(angles_export, normalized_values_list_export, alpha=0.25, color='#ffb3c6')

                            ax_radar_export.set_xticks(angles_export[:-1])
                            short_labels_export = [label.split('(')[0].strip()[:20] for label in indicators_export]
                            ax_radar_export.set_xticklabels(short_labels_export, size=8, color='#4a5568', fontweight='600')

                            ax_radar_export.set_ylim(0, 1)
                            ax_radar_export.set_title('Phân tích Đa chiều các Chỉ số\n(Normalized 0-1)',
                                              fontsize=14, fontweight='bold', color='#c2185b', pad=20)
                            ax_radar_export.grid(True, alpha=0.3, linestyle='--', linewidth=0.8, color='#ff6b9d')
                            ax_radar_export.set_facecolor('#ffffff')
                            plt.tight_layout()

                            # Tạo PD label
                            if pd.notna(probs) and pd.notna(preds):
                                pd_label_text = "Default (Vỡ nợ)" if preds == 1 else "Non-Default (Không vỡ nợ)"
                            else:
                                pd_label_text = "N/A"

                            # Generate Word
                            word_buffer = generate_word_report(
                                ratios_display=ratios_display,
                                pd_value=probs if pd.notna(probs) else np.nan,
                                pd_label=pd_label_text,
                                ai_analysis=ai_analysis_text,
                                fig_bar=fig_bar_export,
                                fig_radar=fig_radar_export,
                                company_name=company_name_input
                            )

                            # Close figures
                            plt.close(fig_bar_export)
                            plt.close(fig_radar_export)

                        st.success("✅ Báo cáo Word đã được tạo thành công!")

                        # Download button
                        st.download_button(
                            label="💾 Tải xuống Báo cáo Word",
                            data=word_buffer,
                            file_name=f"BaoCao_TinDung_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

                    except Exception as e:
                        st.error(f"❌ Lỗi khi tạo Word: {str(e)}")
                        st.exception(e)

    else:
        st.info("Hãy tải **ho_so_dn.xlsx** (đủ 3 sheet) để tính X1…X14, dự báo PD và phân tích AI.")

    # Nút lên đầu trang
    st.markdown("""
        <div style='text-align: center; margin-top: 40px; margin-bottom: 20px;'>
            <a href='#top' onclick='window.scrollTo({top: 0, behavior: "smooth"}); return false;' style='text-decoration: none;'>
                <button style='
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    color: white;
                    border: none;
                    padding: 12px 30px;
                    border-radius: 25px;
                    font-size: 16px;
                    font-weight: 600;
                    cursor: pointer;
                    box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
                    transition: all 0.3s ease;
                '>
                    ⬆️ Lên đầu trang
                </button>
            </a>
        </div>
    """, unsafe_allow_html=True)

# ========================================
# TAB: DASHBOARD TÀI CHÍNH DOANH NGHIỆP (GSO)
# ========================================
with tab_dashboard:
    st.header("📊 Dashboard Tài chính Doanh nghiệp Việt Nam")
    st.markdown("""
    Dashboard này hiển thị các xu hướng tài chính của doanh nghiệp Việt Nam theo quý,
    dựa trên dữ liệu từ **Tổng cục Thống kê (GSO) - General Statistics Office of Vietnam**.
    """)

    st.divider()

    # Khu vực upload và hướng dẫn
    info_container = st.container(border=True)
    with info_container:
        st.markdown("### 📥 Nguồn Dữ liệu")

        # Highlight tính năng mới
        st.success("""
        🆕 **TÍNH NĂNG MỚI**: Tự động lấy dữ liệu tài chính doanh nghiệp Việt Nam từ **Gemini AI**!
        - ✅ Tự động cập nhật theo tháng (cache 30 ngày)
        - ✅ Dữ liệu từ nguồn tin cậy (GSO, Bộ KH&ĐT)
        - ✅ Không cần tải file thủ công
        """)

        with st.expander("📖 Hướng dẫn sử dụng các nguồn dữ liệu"):
            st.markdown("""
            **🚀 Tự động lấy từ Gemini AI (Khuyến nghị):**
            - Nhấn nút **"Bấm để tạo"** để tự động lấy dữ liệu mới nhất
            - Dữ liệu được cache 30 ngày, tự động cập nhật mỗi tháng
            - Nguồn dữ liệu: GSO, Bộ KH&ĐT, báo cáo kinh tế vĩ mô

            **📂 Tải lên dữ liệu GSO thủ công:**
            1. Truy cập: [https://gso.gov.vn](https://gso.gov.vn)
            2. Chọn mục **Số liệu thống kê** → **Doanh nghiệp**
            3. Tải về file Excel/CSV chứa dữ liệu theo quý
            4. Upload file vào đây để phân tích

            **📊 Dùng Thử:**
            - Sử dụng dữ liệu mẫu để khám phá tính năng

            **Định dạng file yêu cầu (khi upload thủ công):**
            - File CSV hoặc Excel (.xlsx)
            - Cột **Quý/Năm** (ví dụ: Q1-2023, Q2-2023...)
            - Cột **Doanh thu** (đơn vị: tỷ đồng)
            - Cột **Tổng tài sản** (đơn vị: tỷ đồng)
            - Các cột khác: Lợi nhuận, Nợ phải trả, VCSH... (tùy chọn)
            """)

    st.divider()

    # Upload file, lấy dữ liệu từ AI, hoặc sử dụng dữ liệu mẫu
    col_ai, col_sample, col_upload = st.columns([1, 1, 2])

    with col_ai:
        st.markdown("#### 🤖 Dữ liệu lấy từ Gemini")
        use_ai_data = st.button("🚀 Bấm để tạo", use_container_width=True, type="primary",
                                help="Tự động lấy dữ liệu tài chính doanh nghiệp VN mới nhất từ Gemini AI")

    with col_sample:
        st.markdown("#### 🎯 Demo Thử Mẫu")
        use_sample = st.button("📊 Dùng Thử", use_container_width=True, type="secondary")

    with col_upload:
        st.markdown("#### 📂 Tải lên Dữ liệu GSO")
        uploaded_gso = st.file_uploader(
            "Chọn file CSV hoặc Excel chứa dữ liệu GSO",
            type=['csv', 'xlsx'],
            key="gso_upload"
        )

    # Biến lưu DataFrame
    gso_data = None

    # Xử lý upload file
    if uploaded_gso is not None:
        try:
            with st.spinner('Đang đọc dữ liệu từ file...'):
                if uploaded_gso.name.endswith('.csv'):
                    gso_data = pd.read_csv(uploaded_gso)
                else:
                    gso_data = pd.read_excel(uploaded_gso)
            st.success(f"✅ Đã tải thành công file: **{uploaded_gso.name}**")
        except Exception as e:
            st.error(f"❌ Lỗi khi đọc file: {e}")

    # Hoặc lấy dữ liệu tự động từ Gemini AI
    elif use_ai_data:
        if not _GEMINI_OK:
            st.error("❌ Thiếu thư viện google-genai. Vui lòng cài đặt: pip install google-genai")
        else:
            api_key = st.secrets.get("GEMINI_API_KEY")
            if api_key:
                with st.spinner('🤖 Đang lấy dữ liệu tài chính từ Gemini AI... (có thể mất 10-20 giây)'):
                    gso_data = get_financial_data_from_ai(api_key)
                    if gso_data is not None and not gso_data.empty:
                        st.success("✅ Đã lấy thành công dữ liệu tài chính doanh nghiệp Việt Nam từ Gemini AI!")
                        st.info("💡 **Dữ liệu được cache 30 ngày** - Sẽ tự động cập nhật vào tháng sau")
                    else:
                        st.warning("⚠️ Không thể lấy dữ liệu từ AI. Vui lòng thử lại hoặc sử dụng dữ liệu mẫu.")
            else:
                st.error("❌ **Lỗi Khóa API**: Không tìm thấy GEMINI_API_KEY trong Streamlit Secrets.")

    # Hoặc sử dụng dữ liệu mẫu
    elif use_sample:
        st.info("📊 Đang sử dụng dữ liệu mẫu từ GSO (Demo)")
        # Tạo dữ liệu mẫu (giả lập dữ liệu thực từ GSO)
        quarters = [
            'Q1-2021', 'Q2-2021', 'Q3-2021', 'Q4-2021',
            'Q1-2022', 'Q2-2022', 'Q3-2022', 'Q4-2022',
            'Q1-2023', 'Q2-2023', 'Q3-2023', 'Q4-2023',
            'Q1-2024', 'Q2-2024', 'Q3-2024'
        ]

        # Dữ liệu giả lập (xu hướng tăng trưởng)
        np.random.seed(42)
        base_revenue = 5000
        base_assets = 8000
        base_profit = 500
        base_debt = 3500

        revenues = [base_revenue + i*150 + np.random.randint(-100, 200) for i in range(len(quarters))]
        assets = [base_assets + i*200 + np.random.randint(-150, 250) for i in range(len(quarters))]
        profits = [base_profit + i*30 + np.random.randint(-50, 80) for i in range(len(quarters))]
        debts = [base_debt + i*80 + np.random.randint(-100, 150) for i in range(len(quarters))]
        equity = [assets[i] - debts[i] for i in range(len(quarters))]

        gso_data = pd.DataFrame({
            'Quý': quarters,
            'Doanh thu (tỷ VNĐ)': revenues,
            'Tổng tài sản (tỷ VNĐ)': assets,
            'Lợi nhuận (tỷ VNĐ)': profits,
            'Nợ phải trả (tỷ VNĐ)': debts,
            'VCSH (tỷ VNĐ)': equity
        })

    # Hiển thị và phân tích dữ liệu nếu có
    if gso_data is not None:
        st.divider()
        st.markdown("### 📈 Dữ liệu và Phân tích")

        # Hiển thị dữ liệu thô
        with st.expander("🔍 Xem Dữ liệu Thô"):
            st.dataframe(gso_data, use_container_width=True)

            # Thống kê mô tả
            st.markdown("#### Thống kê Mô tả")
            st.dataframe(gso_data.describe(), use_container_width=True)

        st.divider()

        # Phần trực quan hóa
        st.markdown("### 📊 Trực Quan Hóa Xu Hướng Tài Chính")

        # Kiểm tra các cột cần thiết
        required_cols = ['Quý', 'Doanh thu (tỷ VNĐ)', 'Tổng tài sản (tỷ VNĐ)']
        missing_cols = [col for col in required_cols if col not in gso_data.columns]

        if missing_cols:
            st.warning(f"⚠️ File dữ liệu thiếu các cột: {', '.join(missing_cols)}. Vui lòng đảm bảo file có đủ các cột yêu cầu.")
        else:
            # Biểu đồ 1: Xu hướng Doanh thu theo quý
            st.markdown("#### 💰 Xu hướng Doanh thu theo Quý")
            fig1, ax1 = plt.subplots(figsize=(14, 6))
            fig1.patch.set_facecolor('#fff5f7')
            ax1.set_facecolor('#ffffff')

            # Vẽ đường xu hướng doanh thu
            ax1.plot(gso_data['Quý'], gso_data['Doanh thu (tỷ VNĐ)'],
                    marker='o', linewidth=3, markersize=8, color='#ff6b9d',
                    label='Doanh thu', linestyle='-', alpha=0.9)

            # Fill area under curve
            ax1.fill_between(gso_data['Quý'], gso_data['Doanh thu (tỷ VNĐ)'],
                            alpha=0.2, color='#ffb3c6')

            # Styling
            ax1.set_xlabel('Quý', fontsize=13, fontweight='600', color='#4a5568')
            ax1.set_ylabel('Doanh thu (tỷ VNĐ)', fontsize=13, fontweight='600', color='#4a5568')
            ax1.set_title('Xu hướng Doanh thu Doanh nghiệp Việt Nam theo Quý',
                         fontsize=16, fontweight='bold', color='#c2185b', pad=20)
            ax1.grid(True, alpha=0.2, linestyle='--', linewidth=0.8, color='#ff6b9d')
            ax1.legend(fontsize=11, frameon=True, shadow=True)

            # Xoay labels trục x
            plt.xticks(rotation=45, ha='right')

            # Remove top and right spines
            ax1.spines['top'].set_visible(False)
            ax1.spines['right'].set_visible(False)
            ax1.spines['left'].set_color('#d0d0d0')
            ax1.spines['bottom'].set_color('#d0d0d0')

            plt.tight_layout()
            st.pyplot(fig1)
            plt.close(fig1)

            st.divider()

            # Biểu đồ 2: So sánh Doanh thu và Tổng tài sản
            st.markdown("#### 🏢 So sánh Doanh thu và Tổng Tài sản")
            fig2, ax2 = plt.subplots(figsize=(14, 6))
            fig2.patch.set_facecolor('#fff5f7')
            ax2.set_facecolor('#ffffff')

            # Vẽ 2 đường xu hướng
            ax2.plot(gso_data['Quý'], gso_data['Doanh thu (tỷ VNĐ)'],
                    marker='o', linewidth=2.5, markersize=7, color='#ff6b9d',
                    label='Doanh thu', linestyle='-', alpha=0.9)

            ax2.plot(gso_data['Quý'], gso_data['Tổng tài sản (tỷ VNĐ)'],
                    marker='s', linewidth=2.5, markersize=7, color='#4a90e2',
                    label='Tổng tài sản', linestyle='-', alpha=0.9)

            # Styling
            ax2.set_xlabel('Quý', fontsize=13, fontweight='600', color='#4a5568')
            ax2.set_ylabel('Giá trị (tỷ VNĐ)', fontsize=13, fontweight='600', color='#4a5568')
            ax2.set_title('So sánh Doanh thu và Tổng Tài sản theo Quý',
                         fontsize=16, fontweight='bold', color='#c2185b', pad=20)
            ax2.grid(True, alpha=0.2, linestyle='--', linewidth=0.8, color='#ff6b9d')
            ax2.legend(fontsize=11, frameon=True, shadow=True, loc='upper left')

            plt.xticks(rotation=45, ha='right')

            ax2.spines['top'].set_visible(False)
            ax2.spines['right'].set_visible(False)
            ax2.spines['left'].set_color('#d0d0d0')
            ax2.spines['bottom'].set_color('#d0d0d0')

            plt.tight_layout()
            st.pyplot(fig2)
            plt.close(fig2)

            st.divider()

            # Biểu đồ 3: Biểu đồ cột so sánh các chỉ số (nếu có đủ cột)
            optional_cols = ['Lợi nhuận (tỷ VNĐ)', 'Nợ phải trả (tỷ VNĐ)', 'VCSH (tỷ VNĐ)']
            available_optional = [col for col in optional_cols if col in gso_data.columns]

            if available_optional:
                st.markdown("#### 📊 Phân tích Chi tiết các Chỉ số Tài chính")

                # Chọn quý để so sánh
                selected_quarters = st.multiselect(
                    "🔍 Chọn các quý để so sánh:",
                    options=gso_data['Quý'].tolist(),
                    default=gso_data['Quý'].tolist()[-4:],  # Mặc định 4 quý gần nhất
                    key="quarter_selector"
                )

                if selected_quarters:
                    filtered_data = gso_data[gso_data['Quý'].isin(selected_quarters)]

                    # Tạo biểu đồ cột nhóm
                    fig3, ax3 = plt.subplots(figsize=(14, 7))
                    fig3.patch.set_facecolor('#fff5f7')
                    ax3.set_facecolor('#ffffff')

                    # Số lượng quý và chỉ số
                    n_quarters = len(selected_quarters)
                    n_indicators = len(available_optional)

                    # Vị trí các cột
                    x = np.arange(n_quarters)
                    width = 0.25  # Độ rộng mỗi cột

                    # Màu sắc cho các chỉ số
                    colors = ['#ff6b9d', '#4a90e2', '#50c878']

                    # Vẽ các cột
                    for i, col in enumerate(available_optional):
                        offset = (i - n_indicators/2 + 0.5) * width
                        ax3.bar(x + offset, filtered_data[col], width,
                               label=col.replace(' (tỷ VNĐ)', ''),
                               color=colors[i % len(colors)], alpha=0.8,
                               edgecolor='white', linewidth=1.5)

                    # Styling
                    ax3.set_xlabel('Quý', fontsize=13, fontweight='600', color='#4a5568')
                    ax3.set_ylabel('Giá trị (tỷ VNĐ)', fontsize=13, fontweight='600', color='#4a5568')
                    ax3.set_title('So sánh các Chỉ số Tài chính theo Quý',
                                 fontsize=16, fontweight='bold', color='#c2185b', pad=20)
                    ax3.set_xticks(x)
                    ax3.set_xticklabels(selected_quarters, rotation=45, ha='right')
                    ax3.legend(fontsize=11, frameon=True, shadow=True)
                    ax3.grid(True, alpha=0.2, linestyle='--', linewidth=0.8, color='#ff6b9d', axis='y')

                    ax3.spines['top'].set_visible(False)
                    ax3.spines['right'].set_visible(False)
                    ax3.spines['left'].set_color('#d0d0d0')
                    ax3.spines['bottom'].set_color('#d0d0d0')

                    plt.tight_layout()
                    st.pyplot(fig3)
                    plt.close(fig3)

            st.divider()

            # Phần thống kê tổng quan
            st.markdown("### 📈 Thống Kê Tổng Quan")

            metric_col1, metric_col2, metric_col3, metric_col4 = st.columns(4)

            with metric_col1:
                avg_revenue = gso_data['Doanh thu (tỷ VNĐ)'].mean()
                st.metric(
                    label="Doanh thu TB",
                    value=f"{avg_revenue:,.0f} tỷ",
                    delta=f"{gso_data['Doanh thu (tỷ VNĐ)'].iloc[-1] - gso_data['Doanh thu (tỷ VNĐ)'].iloc[0]:,.0f} tỷ",
                    delta_color="normal"
                )

            with metric_col2:
                avg_assets = gso_data['Tổng tài sản (tỷ VNĐ)'].mean()
                st.metric(
                    label="Tổng TS TB",
                    value=f"{avg_assets:,.0f} tỷ",
                    delta=f"{gso_data['Tổng tài sản (tỷ VNĐ)'].iloc[-1] - gso_data['Tổng tài sản (tỷ VNĐ)'].iloc[0]:,.0f} tỷ",
                    delta_color="normal"
                )

            if 'Lợi nhuận (tỷ VNĐ)' in gso_data.columns:
                with metric_col3:
                    avg_profit = gso_data['Lợi nhuận (tỷ VNĐ)'].mean()
                    st.metric(
                        label="Lợi nhuận TB",
                        value=f"{avg_profit:,.0f} tỷ",
                        delta=f"{gso_data['Lợi nhuận (tỷ VNĐ)'].iloc[-1] - gso_data['Lợi nhuận (tỷ VNĐ)'].iloc[0]:,.0f} tỷ",
                        delta_color="normal"
                    )

            if 'VCSH (tỷ VNĐ)' in gso_data.columns:
                with metric_col4:
                    avg_equity = gso_data['VCSH (tỷ VNĐ)'].mean()
                    st.metric(
                        label="VCSH TB",
                        value=f"{avg_equity:,.0f} tỷ",
                        delta=f"{gso_data['VCSH (tỷ VNĐ)'].iloc[-1] - gso_data['VCSH (tỷ VNĐ)'].iloc[0]:,.0f} tỷ",
                        delta_color="normal"
                    )

            st.divider()

            # Kết luận và Insights
            st.markdown("### 💡 Nhận xét và Insights")

            insights_container = st.container(border=True)
            with insights_container:
                # Tính toán tốc độ tăng trưởng
                revenue_growth = ((gso_data['Doanh thu (tỷ VNĐ)'].iloc[-1] - gso_data['Doanh thu (tỷ VNĐ)'].iloc[0]) / gso_data['Doanh thu (tỷ VNĐ)'].iloc[0]) * 100
                assets_growth = ((gso_data['Tổng tài sản (tỷ VNĐ)'].iloc[-1] - gso_data['Tổng tài sản (tỷ VNĐ)'].iloc[0]) / gso_data['Tổng tài sản (tỷ VNĐ)'].iloc[0]) * 100

                st.markdown(f"""
                **Xu hướng Tăng trưởng:**
                - 📈 **Doanh thu**: Tăng trưởng **{revenue_growth:.1f}%** từ quý đầu đến quý cuối
                - 🏢 **Tổng tài sản**: Tăng trưởng **{assets_growth:.1f}%** từ quý đầu đến quý cuối

                **Đánh giá:**
                """)

                if revenue_growth > 10:
                    st.success("✅ Doanh nghiệp có xu hướng tăng trưởng doanh thu tốt (>10%)")
                elif revenue_growth > 0:
                    st.info("💡 Doanh nghiệp có tăng trưởng doanh thu nhẹ")
                else:
                    st.warning("⚠️ Doanh nghiệp có xu hướng giảm doanh thu, cần xem xét kỹ")

                if assets_growth > 15:
                    st.success("✅ Quy mô tài sản tăng trưởng mạnh (>15%)")
                elif assets_growth > 0:
                    st.info("💡 Quy mô tài sản có tăng trưởng")
                else:
                    st.warning("⚠️ Quy mô tài sản giảm, cần phân tích nguyên nhân")

    else:
        st.info("💡 Vui lòng tải lên file dữ liệu GSO hoặc sử dụng dữ liệu mẫu để xem phân tích.")

    # Nút lên đầu trang
    st.markdown("""
        <div style='text-align: center; margin-top: 40px; margin-bottom: 20px;'>
            <a href='#top' onclick='window.scrollTo({top: 0, behavior: "smooth"}); return false;' style='text-decoration: none;'>
                <button style='
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    color: white;
                    border: none;
                    padding: 12px 30px;
                    border-radius: 25px;
                    font-size: 16px;
                    font-weight: 600;
                    cursor: pointer;
                    box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
                    transition: all 0.3s ease;
                '>
                    ⬆️ Lên đầu trang
                </button>
            </a>
        </div>
    """, unsafe_allow_html=True)

# ========================================
# TAB: TIN TỨC TÀI CHÍNH
# ========================================
with tab_news:
    st.header("📰 Tin tức Tài chính")
    st.markdown("""
    Tin tức tài chính mới nhất từ các nguồn uy tín tại Việt Nam.
    Dữ liệu tự động cập nhật mỗi **120 phút**.
    """)

    st.divider()

    if not _FEEDPARSER_OK:
        st.error("⚠️ **Thiếu thư viện feedparser**. Vui lòng cài đặt: `pip install feedparser python-dateutil`")
    else:
        # Định nghĩa các nguồn RSS
        rss_sources = {
            "📊 CafeF": "https://cafef.vn/thi-truong-chung-khoan.rss",
            "💼 Vietstock": "https://vietstock.vn/rss/tai-chinh.rss",
            "💰 Báo Đầu tư": "https://baodautu.vn/rss/kinh-doanh.rss",
            "🏢 VNExpress Kinh doanh": "https://vnexpress.net/rss/kinh-doanh.rss"
        }

        # Hiển thị thời gian cập nhật
        col_update, col_cache = st.columns([3, 1])
        with col_update:
            st.caption(f"🕐 Cập nhật: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        with col_cache:
            st.caption("♻️ Cache: 120 phút")

        st.divider()

        # Tạo layout 2 cột
        col1, col2 = st.columns(2)

        sources_list = list(rss_sources.items())

        # Hiển thị nguồn tin 1 và 2 ở cột trái
        with col1:
            # Nguồn 1: CafeF
            source_name, source_url = sources_list[0]
            with st.container(border=True):
                st.markdown(f"### {source_name}")
                articles = fetch_rss_feed(source_url, source_name)

                for i, article in enumerate(articles):
                    st.markdown(f"""
                    <div style='
                        padding: 10px;
                        margin: 8px 0;
                        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
                        border-radius: 8px;
                        border-left: 4px solid #667eea;
                    '>
                        <div style='font-size: 14px; font-weight: 600; color: #2c3e50; margin-bottom: 5px;'>
                            📌 {article['title']}
                        </div>
                        <div style='font-size: 12px; color: #7f8c8d; margin-bottom: 8px;'>
                            🕐 {article['published']}
                        </div>
                        <a href='{article['link']}' target='_blank' style='
                            color: #667eea;
                            text-decoration: none;
                            font-size: 12px;
                            font-weight: 600;
                        '>
                            🔗 Đọc chi tiết →
                        </a>
                    </div>
                    """, unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)

            # Nguồn 3: Báo Đầu tư
            source_name, source_url = sources_list[2]
            with st.container(border=True):
                st.markdown(f"### {source_name}")
                articles = fetch_rss_feed(source_url, source_name)

                for i, article in enumerate(articles):
                    st.markdown(f"""
                    <div style='
                        padding: 10px;
                        margin: 8px 0;
                        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
                        border-radius: 8px;
                        border-left: 4px solid #667eea;
                    '>
                        <div style='font-size: 14px; font-weight: 600; color: #2c3e50; margin-bottom: 5px;'>
                            📌 {article['title']}
                        </div>
                        <div style='font-size: 12px; color: #7f8c8d; margin-bottom: 8px;'>
                            🕐 {article['published']}
                        </div>
                        <a href='{article['link']}' target='_blank' style='
                            color: #667eea;
                            text-decoration: none;
                            font-size: 12px;
                            font-weight: 600;
                        '>
                            🔗 Đọc chi tiết →
                        </a>
                    </div>
                    """, unsafe_allow_html=True)

        # Hiển thị nguồn tin 2 và 4 ở cột phải
        with col2:
            # Nguồn 2: Vietstock
            source_name, source_url = sources_list[1]
            with st.container(border=True):
                st.markdown(f"### {source_name}")
                articles = fetch_rss_feed(source_url, source_name)

                for i, article in enumerate(articles):
                    st.markdown(f"""
                    <div style='
                        padding: 10px;
                        margin: 8px 0;
                        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
                        border-radius: 8px;
                        border-left: 4px solid #667eea;
                    '>
                        <div style='font-size: 14px; font-weight: 600; color: #2c3e50; margin-bottom: 5px;'>
                            📌 {article['title']}
                        </div>
                        <div style='font-size: 12px; color: #7f8c8d; margin-bottom: 8px;'>
                            🕐 {article['published']}
                        </div>
                        <a href='{article['link']}' target='_blank' style='
                            color: #667eea;
                            text-decoration: none;
                            font-size: 12px;
                            font-weight: 600;
                        '>
                            🔗 Đọc chi tiết →
                        </a>
                    </div>
                    """, unsafe_allow_html=True)

            st.markdown("<br>", unsafe_allow_html=True)

            # Nguồn 4: VNExpress
            source_name, source_url = sources_list[3]
            with st.container(border=True):
                st.markdown(f"### {source_name}")
                articles = fetch_rss_feed(source_url, source_name)

                for i, article in enumerate(articles):
                    st.markdown(f"""
                    <div style='
                        padding: 10px;
                        margin: 8px 0;
                        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
                        border-radius: 8px;
                        border-left: 4px solid #667eea;
                    '>
                        <div style='font-size: 14px; font-weight: 600; color: #2c3e50; margin-bottom: 5px;'>
                            📌 {article['title']}
                        </div>
                        <div style='font-size: 12px; color: #7f8c8d; margin-bottom: 8px;'>
                            🕐 {article['published']}
                        </div>
                        <a href='{article['link']}' target='_blank' style='
                            color: #667eea;
                            text-decoration: none;
                            font-size: 12px;
                            font-weight: 600;
                        '>
                            🔗 Đọc chi tiết →
                        </a>
                    </div>
                    """, unsafe_allow_html=True)

    # Nút lên đầu trang
    st.markdown("""
        <div style='text-align: center; margin-top: 40px; margin-bottom: 20px;'>
            <a href='#top' onclick='window.scrollTo({top: 0, behavior: "smooth"}); return false;' style='text-decoration: none;'>
                <button style='
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    color: white;
                    border: none;
                    padding: 12px 30px;
                    border-radius: 25px;
                    font-size: 16px;
                    font-weight: 600;
                    cursor: pointer;
                    box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
                    transition: all 0.3s ease;
                '>
                    ⬆️ Lên đầu trang
                </button>
            </a>
        </div>
    """, unsafe_allow_html=True)

# ========================================
# TAB: NHÓM TÁC GIẢ
# ========================================
with tab_authors:
    # Header với hiệu ứng gradient
    st.markdown("""
        <div style='text-align: center; padding: 30px; background: linear-gradient(135deg, #fbc2eb 0%, #a6c1ee 100%); border-radius: 15px; margin-bottom: 30px; box-shadow: 0 10px 30px rgba(102, 126, 234, 0.3);'>
            <h1 style='color: white; margin: 0; font-size: 2.5rem; font-weight: 700; text-shadow: 2px 2px 4px rgba(0,0,0,0.2);'>
                👥 NHÓM ÁNH SÁNG SỐ
            </h1>
            <p style='color: #f0f0f0; font-size: 1.1rem; margin-top: 10px; font-weight: 500;'>
                Cuộc thi Agribank làm chủ công nghệ trong kỷ nguyên số 2025
            </p>
        </div>
    """, unsafe_allow_html=True)

    # Ảnh nhóm ở giữa
    col_left, col_center, col_right = st.columns([1, 2, 1])
    with col_center:
        try:
            st.image("NHOM ANH SANG SO.jpg", use_container_width=True, caption="Team Ánh Sáng Số - Ánh sáng của đổi mới, bước đi của tương lai")
        except:
            st.info("📸 Ảnh nhóm: NHOM ANH SANG SO.jpg")

    st.markdown("<br>", unsafe_allow_html=True)

    # Giới thiệu chung
    st.markdown("""
        <div style='text-align: center; padding: 20px; background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%); border-radius: 10px; margin-bottom: 40px;'>
            <p style='color: #2c3e50; font-size: 1.1rem; line-height: 1.8; margin: 0;'>
                🌟 Chúng tôi là những Agribanker - những người giữ ánh sáng của niềm tin, lan tỏa tinh thần chuyển đổi số trên mọi miền đất nước
            </p>
        </div>
    """, unsafe_allow_html=True)

    st.markdown("### 🌟 Thành viên nhóm")
    st.markdown("<br>", unsafe_allow_html=True)

    # Profile thành viên 1: Trần Ngọc Trúc Huỳnh
    col1, col2 = st.columns([1, 2])

    with col1:
        try:
            st.image("Tran Ngoc Truc Huynh.jpg", use_container_width=True)
        except:
            st.info("📸 Tran Ngoc Truc Huynh.jpg")

    with col2:
        st.markdown("""
            <div style='background: linear-gradient(135deg, #ffecd2 0%, #fcb69f 100%); padding: 25px; border-radius: 15px; box-shadow: 0 8px 20px rgba(252, 182, 159, 0.3); height: 100%;'>
                <h3 style='color: #d63447; margin-top: 0; font-size: 1.8rem; border-bottom: 3px solid #d63447; padding-bottom: 10px;'>
                    🎯 Trần Ngọc Trúc Huỳnh
                </h3>
                <p style='color: #2c3e50; margin: 15px 0; font-size: 1.05rem;'>
                    <strong>🏢 Chức vụ:</strong> Giao dịch viên<br>
                    <strong>📍 Đơn vị:</strong> Agribank chi nhánh Tiền Giang
                </p>
                <div style='background: rgba(255, 255, 255, 0.6); padding: 15px; border-radius: 10px; margin-top: 15px;'>
                    <p style='color: #d63447; font-weight: 700; margin-bottom: 10px; font-size: 1.1rem;'>💼 Vai trò trong nhóm:</p>
                    <ul style='color: #2c3e50; margin: 0; padding-left: 20px; line-height: 1.8;'>
                        <li>Ý tưởng nâng cấp chương trình phiên bản 2.0</li>
                        <li>Kỹ thuật chính – Coder chính cho mô hình 2.0</li>
                        <li>Trailer giới thiệu mô hình nâng cấp</li>
                        <li>Phân chia, tổ chức công việc nhóm</li>
                        <li>Hỗ trợ kỹ thuật cho Version 1.0</li>
                        <li>Kịch bản & Thuyết trình Demo Version 1.0</li>
                    </ul>
                </div>
            </div>
        """, unsafe_allow_html=True)

    st.markdown("<br><br>", unsafe_allow_html=True)

    # Profile thành viên 2: Nguyễn Hồng Cường
    col1, col2 = st.columns([1, 2])

    with col1:
        try:
            st.image("NGUYEN HONG CUONG.jpg", use_container_width=True)
        except:
            st.info("📸 NGUYEN HONG CUONG.jpg")

    with col2:
        st.markdown("""
            <div style='background: linear-gradient(135deg, #a1c4fd 0%, #c2e9fb 100%); padding: 25px; border-radius: 15px; box-shadow: 0 8px 20px rgba(161, 196, 253, 0.3); height: 100%;'>
                <h3 style='color: #2c5aa0; margin-top: 0; font-size: 1.8rem; border-bottom: 3px solid #2c5aa0; padding-bottom: 10px;'>
                    🎯 Nguyễn Hồng Cường
                </h3>
                <p style='color: #2c3e50; margin: 15px 0; font-size: 1.05rem;'>
                    <strong>🏢 Chức vụ:</strong> Trưởng phòng Kiểm tra – Kiểm soát Nội bộ<br>
                    <strong>📍 Đơn vị:</strong> Agribank chi nhánh Đông Hải Phòng
                </p>
                <div style='background: rgba(255, 255, 255, 0.6); padding: 15px; border-radius: 10px; margin-top: 15px;'>
                    <p style='color: #2c5aa0; font-weight: 700; margin-bottom: 10px; font-size: 1.1rem;'>💼 Vai trò trong nhóm:</p>
                    <ul style='color: #2c3e50; margin: 0; padding-left: 20px; line-height: 1.8;'>
                        <li>Kỹ thuật chính – Coder chính mô hình Version 1.0</li>
                        <li>Demo trực tiếp mô hình Version 1.0 trên sân khấu</li>
                        <li>Hỗ trợ kỹ thuật cho mô hình nâng cấp Version 2.0</li>
                    </ul>
                </div>
            </div>
        """, unsafe_allow_html=True)

    st.markdown("<br><br>", unsafe_allow_html=True)

    # Profile thành viên 3: Nguyễn Trung Thành
    col1, col2 = st.columns([1, 2])

    with col1:
        try:
            st.image("NGUYEN TRUNG THANH.jpg", use_container_width=True)
        except:
            st.info("📸 NGUYEN TRUNG THANH.jpg")

    with col2:
        st.markdown("""
            <div style='background: linear-gradient(135deg, #ffeaa7 0%, #fdcb6e 100%); padding: 25px; border-radius: 15px; box-shadow: 0 8px 20px rgba(253, 203, 110, 0.3); height: 100%;'>
                <h3 style='color: #e17055; margin-top: 0; font-size: 1.8rem; border-bottom: 3px solid #e17055; padding-bottom: 10px;'>
                    🎯 Nguyễn Trung Thành
                </h3>
                <p style='color: #2c3e50; margin: 15px 0; font-size: 1.05rem;'>
                    <strong>🏢 Chức vụ:</strong> Phó trưởng Phòng Kế toán Ngân quỹ<br>
                    <strong>📍 Đơn vị:</strong> Agribank chi nhánh Hải Dương
                </p>
                <div style='background: rgba(255, 255, 255, 0.6); padding: 15px; border-radius: 10px; margin-top: 15px;'>
                    <p style='color: #e17055; font-weight: 700; margin-bottom: 10px; font-size: 1.1rem;'>💼 Vai trò trong nhóm:</p>
                    <ul style='color: #2c3e50; margin: 0; padding-left: 20px; line-height: 1.8;'>
                        <li>Hỗ trợ kỹ thuật cho mô hình Version 1.0</li>
                        <li>Thuyết trình sân khấu Demo Version 1.0</li>
                        <li>Thiết kế Poster mô hình Version 1.0</li>
                    </ul>
                </div>
            </div>
        """, unsafe_allow_html=True)

    st.markdown("<br><br>", unsafe_allow_html=True)

    # Thông điệp kết thúc
    st.markdown("""
        <div style='text-align: center; padding: 30px; background: linear-gradient(135deg, #ff6b9d 0%, #c06c84 100%); border-radius: 15px; margin-top: 40px; box-shadow: 0 10px 30px rgba(255, 107, 157, 0.3);'>
            <h3 style='color: white; margin: 0 0 15px 0; font-size: 1.8rem;'>🚀 Sứ mệnh của chúng tôi</h3>
            <p style='color: #fff; font-size: 1.1rem; line-height: 1.8; margin: 0;'>
                Ứng dụng trí tuệ nhân tạo và công nghệ số để nâng cao hiệu quả hoạt động, quản trị rủi ro và chất lượng phục vụ khách hàng, góp phần hiện thực hóa chiến lược chuyển đổi số của Agribank.
            </p>
            <div style='margin-top: 20px; font-size: 2rem;'>
                💡 🎯 🌟 💼 🏆
            </div>
        </div>
    """, unsafe_allow_html=True)

    # Nút lên đầu trang
    st.markdown("""
        <div style='text-align: center; margin-top: 40px; margin-bottom: 20px;'>
            <a href='#top' onclick='window.scrollTo({top: 0, behavior: "smooth"}); return false;' style='text-decoration: none;'>
                <button style='
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    color: white;
                    border: none;
                    padding: 12px 30px;
                    border-radius: 25px;
                    font-size: 16px;
                    font-weight: 600;
                    cursor: pointer;
                    box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
                    transition: all 0.3s ease;
                '>
                    ⬆️ Lên đầu trang
                </button>
            </a>
        </div>
    """, unsafe_allow_html=True)

# ========================================
# PREMIUM BANKING FOOTER
# ========================================
st.markdown("---")
footer_col1, footer_col2, footer_col3 = st.columns([2, 2, 1])

with footer_col1:
    st.markdown("""
    <div style='padding: 15px; text-align: left;'>
        <h4 style='color: #ff6b9d; margin-bottom: 10px;'>🏦 Chương Trình Đánh Giá Rủi Ro Tín Dụng</h4>
        <p style='color: #6b7280; font-size: 0.9rem; margin: 5px 0;'>
            Giải pháp AI tiên tiến cho phân tích tài chính doanh nghiệp
        </p>
        <p style='color: #6b7280; font-size: 0.85rem; margin: 5px 0;'>
            Authored by <strong>ÁNH SÁNG SỐ Team</strong> 
        </p>
    </div>
    """, unsafe_allow_html=True)

with footer_col2:
    st.markdown("""
    <div style='padding: 15px; text-align: left;'>
        <h4 style='color: #ff6b9d; margin-bottom: 10px;'>📊 Tính Năng Chính</h4>
        <ul style='color: #6b7280; font-size: 0.85rem; margin: 5px 0; padding-left: 20px;'>
            <li>Phân tích 14 chỉ số tài chính tự động</li>
            <li>Dự báo xác suất vỡ nợ (PD) và Phân tích chuyên sâu</li>
            <li>DashBoard Tài Chính Doanh Nghiệp tổng quan</li>
            <li>Tin tức tài chính cập nhật Real-Time</li>
        </ul>
    </div>
    """, unsafe_allow_html=True)

with footer_col3:
    st.markdown(f"""
    <div style='padding: 15px; text-align: center;'>
        <div style='font-size: 3rem; margin-bottom: 10px;'>💖</div>
        <p style='color: #ffb3c6; font-weight: 700; font-size: 0.9rem; margin: 5px 0;'>
            SWEET ANALYTICS
        </p>
        <p style='color: #6b7280; font-size: 0.75rem;'>
            Version 2.0 Premium
        </p>
    </div>
    """, unsafe_allow_html=True)

st.markdown(f"""
<div style='text-align: center; padding: 20px; margin-top: 20px;
            background: linear-gradient(135deg, #ff6b9d 0%, #ff85a1 100%);
            border-radius: 15px; box-shadow: 0 4px 15px rgba(255, 107, 157, 0.2);'>
    <p style='color: #ffffff; margin: 5px 0; font-size: 0.9rem; font-weight: 600;'>
        © {datetime.now().year} Credit Risk Assessment System | Developed with ❤️ using Streamlit
    </p>
    <p style='color: #fff0f5; margin: 5px 0; font-size: 0.85rem;'>
        🔒 Secure • 🚀 Fast • 🎯 Accurate • ✨ AI-Powered
    </p>
</div>
""", unsafe_allow_html=True)
